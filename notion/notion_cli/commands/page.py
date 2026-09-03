"""Page commands for Notion CLI - manage standalone pages not in databases."""
import json
import typer
from typing import Optional, List
from pathlib import Path

COMMAND_CREDENTIALS = {
    "search": ["custom"],
    "list": ["custom"],
    "get": ["custom"],
    "create": ["custom"],
    "import": ["custom"],
    "export": ["custom"],
    "duplicate": ["custom"],
    "update": ["custom"],
    "delete": ["custom"],
    "content": ["custom"],
    "blocks": ["custom"],
    "files": ["custom"],
}

from ..client import flatten_block_tree, get_client
from ..block_limits import enforce_block_limits, find_oversize_rich_text as _find_oversize_rich_text
from ..icons import IconFormatError, parse_icon
from ..markdown_images import process_markdown_images
from ..downloads import download_files
from cli_tools_shared import confirm_destructive_action
from cli_tools_shared.filters import validate_filters, apply_filters, FilterValidationError
from ..output import (
    print_json,
    print_table,
    command,
    format_page_for_display,
    format_block_for_display,
    build_text_block_update,
    TEXT_EDITABLE_BLOCK_TYPES,
    print_error,
    print_success,
    print_warning,
    blocks_to_markdown,
    text_to_blocks,
    block_to_markdown,
)

app = typer.Typer(help="Query and manage standalone pages")
content_app = typer.Typer(help="Manage page content (blocks)")
blocks_app = typer.Typer(help="Manage individual blocks")
files_app = typer.Typer(help="Download file attachments from pages")
app.add_typer(content_app, name="content")
app.add_typer(blocks_app, name="blocks")
app.add_typer(files_app, name="files")


def _page_file_downloads(blocks: List[dict]) -> List[dict]:
    """Return download metadata for file blocks from the public API shape."""
    downloads = []
    for block in flatten_block_tree(blocks):
        if block.get("type") != "file":
            continue

        payload = block.get("file", {})
        source_type = payload.get("type")
        source = payload.get(source_type, {}) if source_type else {}
        name = payload.get("name")
        url = source.get("url")
        if not name or not url:
            raise ValueError(
                f"File block {block.get('id', '<unknown>')} is missing its name or download URL"
            )

        downloads.append(
            {
                "block_id": block.get("id", ""),
                "name": name,
                "source_type": source_type,
                "url": url,
                "expiry_time": source.get("expiry_time"),
            }
        )
    return downloads


@files_app.command("download")
@command
def page_files_download(
    page_id: str = typer.Argument(..., help="Page ID containing file blocks"),
    output: str = typer.Option(
        ...,
        "--output",
        "-o",
        help="Directory where file attachments will be written",
    ),
    table: bool = typer.Option(False, "--table", "-t", help="Display result as a table"),
    force: bool = typer.Option(
        False,
        "--force",
        "-F",
        help="Overwrite files that already exist",
    ),
):
    """Download every file block on an accessible Notion page.

    Notion-hosted file URLs are refreshed by reading the page immediately
    before download because the API signs them for one hour.

    Examples:
        notion pages files download PAGE_ID --output ./skills
        notion pages files download PAGE_ID -o ./skills --table
    """
    client = get_client()
    blocks = client.get_block_children_all(page_id, recursive=True)
    files = _page_file_downloads(blocks)

    results = download_files(files, output, force=force)

    if table:
        print_table(
            results,
            columns=["name", "source_type", "bytes", "output"],
            headers=["Name", "Source", "Bytes", "Output"],
        )
    else:
        print_json(results)

    print_success(f"Downloaded {len(results)} file(s) to {output}")


def format_search_result_for_display(result: dict) -> dict:
    """
    Format a search result (page or database) for display.

    Args:
        result: Raw result from Notion search API

    Returns:
        Simplified record for display
    """
    obj_type = result.get("object", "")

    # Extract title based on object type
    title = ""
    if obj_type == "page":
        props = result.get("properties", {})
        title_prop = props.get("title", {})
        if title_prop.get("type") == "title":
            title_arr = title_prop.get("title", [])
            title = "".join(t.get("plain_text", "") for t in title_arr)
    elif obj_type == "database":
        title_arr = result.get("title", [])
        title = "".join(t.get("plain_text", "") for t in title_arr)

    # Get parent info
    parent = result.get("parent", {})
    parent_type = parent.get("type", "")
    parent_id = ""
    if parent_type == "page_id":
        parent_id = parent.get("page_id", "")
    elif parent_type == "database_id":
        parent_id = parent.get("database_id", "")
    elif parent_type == "workspace":
        parent_id = "workspace"

    return {
        "id": result.get("id", ""),
        "type": obj_type,
        "title": title,
        "parent_type": parent_type,
        "parent_id": parent_id,
        "url": result.get("url", ""),
        "last_edited": result.get("last_edited_time", ""),
    }


@app.command("search")
@command
def page_search(
    query: str = typer.Argument(
        ...,
        help="Text to search for in page titles",
    ),
    table: bool = typer.Option(
        False,
        "--table",
        "-t",
        help="Display as formatted table",
    ),
    sort: Optional[str] = typer.Option(
        None,
        "--sort",
        help="Sort direction by last edited time (asc/desc)",
    ),
    limit: int = typer.Option(
        100,
        "--limit",
        "-l",
        help="Maximum number of results to return",
    ),
):
    """
    Search for pages by title.

    Examples:
        notion pages search "meeting notes"
        notion pages search "project" --table
        notion pages search "draft" --sort desc --limit 10
    """
    client = get_client()

    # Parse sort direction
    sort_direction = None
    if sort:
        if sort.lower() in ("desc", "descending"):
            sort_direction = "descending"
        elif sort.lower() in ("asc", "ascending"):
            sort_direction = "ascending"
        else:
            print_warning(f"Invalid sort value: {sort}. Use 'asc' or 'desc'.")
            raise typer.Exit(1)

    # Search for pages only
    results = client.search_all(
        query=query,
        filter_type="page",
        sort_direction=sort_direction,
        limit=limit,
    )

    # Format results
    formatted = [format_search_result_for_display(r) for r in results]

    if table:
        print_table(
            formatted,
            columns=["title", "parent_type", "last_edited"],
            headers=["Title", "Parent Type", "Last Edited"],
        )
    else:
        print_json(formatted)

    typer.echo(f"\n{len(formatted)} page(s) found.", err=True)


@app.command("list")
@command
def page_list(
    table: bool = typer.Option(
        False,
        "--table",
        "-t",
        help="Display as formatted table",
    ),
    sort: Optional[str] = typer.Option(
        None,
        "--sort",
        help="Sort direction by last edited time (asc/desc)",
    ),
    limit: int = typer.Option(
        100,
        "--limit",
        "-l",
        help="Maximum number of results to return",
    ),
    filter: Optional[List[str]] = typer.Option(
        None,
        "--filter",
        "-f",
        help="Filter using field:op:value format (e.g., title:like:%meeting%). Client-side filtering.",
    ),
    properties: Optional[str] = typer.Option(
        None,
        "--properties",
        "-p",
        help="Comma-separated fields to display (id,type,title,parent_type,parent_id,url,last_edited)",
    ),
):
    """
    List all accessible pages.

    Examples:
        notion pages list
        notion pages list --table
        notion pages list --sort desc --limit 20
        notion pages list --filter "title:like:%meeting%"
        notion pages list --filter "type:eq:page" --filter "parent_type:eq:workspace"
        notion pages list --properties "title,last_edited" --table

    Note: Filtering is performed client-side as the Notion search API has limited filter support.
    """
    client = get_client()

    # Validate filters
    if filter:
        try:
            validate_filters(filter)
        except FilterValidationError as e:
            print_warning(str(e))
            raise typer.Exit(1)

    # Parse sort direction
    sort_direction = None
    if sort:
        if sort.lower() in ("desc", "descending"):
            sort_direction = "descending"
        elif sort.lower() in ("asc", "ascending"):
            sort_direction = "ascending"
        else:
            print_warning(f"Invalid sort value: {sort}. Use 'asc' or 'desc'.")
            raise typer.Exit(1)

    # List all pages (search with no query)
    results = client.search_all(
        query=None,
        filter_type="page",
        sort_direction=sort_direction,
        limit=limit,
    )

    # Format results
    formatted = [format_search_result_for_display(r) for r in results]

    # Apply client-side filter if provided (Notion search API doesn't support rich filtering)
    if filter:
        formatted = apply_filters(formatted, filter)

    # Parse properties option
    display_columns = None
    display_headers = None
    if properties:
        prop_list = [p.strip() for p in properties.split(",")]
        display_columns = prop_list
        # Generate headers by capitalizing and replacing underscores
        display_headers = [p.replace("_", " ").title() for p in prop_list]

    if table:
        cols = display_columns or ["title", "parent_type", "last_edited"]
        hdrs = display_headers or ["Title", "Parent Type", "Last Edited"]
        print_table(formatted, columns=cols, headers=hdrs)
    else:
        # Filter properties for JSON output
        if display_columns:
            formatted = [{k: v for k, v in r.items() if k in display_columns} for r in formatted]
        print_json(formatted)

    typer.echo(f"\n{len(formatted)} page(s) found.", err=True)


@app.command("get")
@command
def page_get(
    page_id: str = typer.Argument(
        ...,
        help="The page ID to retrieve",
    ),
    table: bool = typer.Option(
        False,
        "--table",
        "-t",
        help="Display as formatted table",
    ),
    include_blocks: bool = typer.Option(
        False,
        "--include-blocks",
        "-b",
        help="Include page content blocks in output",
    ),
    markdown: bool = typer.Option(
        False,
        "--markdown",
        "-m",
        help="Output blocks as markdown (requires --include-blocks)",
    ),
    out_file: Optional[str] = typer.Option(
        None,
        "--out-file",
        "-o",
        help="Write markdown content to file (requires --include-blocks --markdown)",
    ),
):
    """
    Get a specific page by ID.

    Examples:
        notion pages get abc123-def456
        notion pages get abc123-def456 --table
        notion pages get abc123-def456 --include-blocks
        notion pages get abc123-def456 --include-blocks --markdown
        notion pages get abc123-def456 -b -m --out-file content.md
    """
    if out_file and not (include_blocks and markdown):
        raise ValueError("--out-file requires --include-blocks and --markdown")

    client = get_client()
    page = client.get_page(page_id)

    formatted = format_page_for_display(page)

    # Fetch and include blocks if requested
    if include_blocks:
        blocks = client.get_block_children_all(page_id, recursive=True)
        if markdown:
            markdown_content = blocks_to_markdown(blocks)

            # Write to file if requested
            if out_file:
                with open(out_file, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                print_success(f"Markdown content written to {out_file}")
                return

            # Output raw markdown (not JSON)
            typer.echo(markdown_content)
            return
        else:
            formatted["blocks"] = blocks

    if table:
        rows = [{"field": k, "value": str(v)} for k, v in formatted.items()]
        print_table(rows, ["field", "value"], ["Field", "Value"])
    else:
        print_json(formatted)


@app.command("create")
@command
def page_create(
    parent_page_id: str = typer.Argument(
        ...,
        help="The parent page ID to create the page under",
    ),
    title: str = typer.Option(
        ...,
        "--title",
        "-t",
        help="Page title",
    ),
    content_file: Optional[str] = typer.Option(
        None,
        "--content-file",
        "-f",
        help="File containing markdown content for the page body",
    ),
    icon: Optional[str] = typer.Option(
        None,
        "--icon",
        help=(
            "Page icon: 'emoji:<shortcode>' (e.g. 'emoji:rocket'), "
            "a literal emoji character (e.g. '🚀'), or 'url:https://...'"
        ),
    ),
    is_toggleable: bool = typer.Option(
        False,
        "--is-toggleable",
        help="Make all heading_1/2/3 blocks generated from --content-file into toggle headings",
    ),
):
    """
    Create a new page under an existing parent page.

    Examples:
        notion pages create PARENT_PAGE_ID --title "My New Page"
        notion pages create PARENT_PAGE_ID -t "Notes" --content-file notes.md
        notion pages create PARENT_PAGE_ID -t "Project" --icon "emoji:rocket"
        notion pages create PARENT_PAGE_ID -t "Project" --icon "🚀"
        notion pages create PARENT_PAGE_ID -t "Outline" --content-file outline.md --is-toggleable
    """
    try:
        client = get_client()

        # Parse icon
        icon_obj = None
        if icon:
            try:
                icon_obj = parse_icon(icon)
            except IconFormatError as exc:
                print_error(str(exc))
                raise typer.Exit(1)

        # Load content file if provided
        children = None
        if content_file:
            with open(content_file, 'r', encoding='utf-8') as f:
                content = f.read()
            # Upload local images BEFORE the page is created so a missing file
            # fails without leaving a half-populated page behind.
            image_uploads = process_markdown_images(content, content_file, client)
            children = text_to_blocks(
                content, image_uploads=image_uploads, is_toggleable=is_toggleable
            )

        # Create the page
        created_page = client.create_standalone_page(
            parent_page_id=parent_page_id,
            title=title,
            children=children,
            icon=icon_obj,
        )

        formatted = format_page_for_display(created_page)
        print_json(formatted)
        print_success(f"Page created successfully: {created_page.get('url', created_page.get('id'))}")

    except FileNotFoundError:
        print_warning(f"File not found: {content_file}")
        raise typer.Exit(1)


def docx_to_notion_blocks(docx_path: str) -> list:
    """
    Convert a Word document to Notion blocks.

    Args:
        docx_path: Path to the .docx file

    Returns:
        List of Notion block objects
    """
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document(docx_path)
    blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()

        # Preserve empty paragraphs as blank lines
        if not text:
            blocks.append({
                "object": "block",
                "type": "paragraph",
                "paragraph": {
                    "rich_text": []
                }
            })
            continue

        style_name = para.style.name.lower() if para.style else ""

        # Determine block type based on style
        if "heading 1" in style_name:
            blocks.append({
                "object": "block",
                "type": "heading_1",
                "heading_1": {
                    "rich_text": [{"type": "text", "text": {"content": text}}]
                }
            })
        elif "heading 2" in style_name:
            blocks.append({
                "object": "block",
                "type": "heading_2",
                "heading_2": {
                    "rich_text": [{"type": "text", "text": {"content": text}}]
                }
            })
        elif "heading 3" in style_name or "heading 4" in style_name:
            blocks.append({
                "object": "block",
                "type": "heading_3",
                "heading_3": {
                    "rich_text": [{"type": "text", "text": {"content": text}}]
                }
            })
        elif "list bullet" in style_name or "list paragraph" in style_name:
            # Check if it looks like a numbered list (starts with number + period)
            import re
            if re.match(r'^\d+\.?\s', text):
                # Strip the number prefix
                text = re.sub(r'^\d+\.?\s*', '', text)
                blocks.append({
                    "object": "block",
                    "type": "numbered_list_item",
                    "numbered_list_item": {
                        "rich_text": [{"type": "text", "text": {"content": text}}]
                    }
                })
            else:
                blocks.append({
                    "object": "block",
                    "type": "bulleted_list_item",
                    "bulleted_list_item": {
                        "rich_text": [{"type": "text", "text": {"content": text}}]
                    }
                })
        elif "quote" in style_name:
            blocks.append({
                "object": "block",
                "type": "quote",
                "quote": {
                    "rich_text": [{"type": "text", "text": {"content": text}}]
                }
            })
        else:
            # Default to paragraph with rich text formatting
            rich_text = []
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue

                annotations = {
                    "bold": run.bold or False,
                    "italic": run.italic or False,
                    "underline": run.underline or False,
                    "strikethrough": run.font.strike or False,
                    "code": "code" in (run.style.name.lower() if run.style else ""),
                }
                # Clean up None values
                annotations = {k: bool(v) for k, v in annotations.items()}

                rich_text.append({
                    "type": "text",
                    "text": {"content": run_text},
                    "annotations": annotations
                })

            # If no runs or empty rich_text, use plain text
            if not rich_text:
                rich_text = [{"type": "text", "text": {"content": text}}]

            blocks.append({
                "object": "block",
                "type": "paragraph",
                "paragraph": {
                    "rich_text": rich_text
                }
            })

    # Handle tables
    for table in doc.tables:
        # Convert table to Notion table block
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells.append([[{"type": "text", "text": {"content": cell_text}}]])
            rows.append(cells)

        if rows:
            # Create table block
            table_block = {
                "object": "block",
                "type": "table",
                "table": {
                    "table_width": len(rows[0]) if rows else 1,
                    "has_column_header": True,
                    "has_row_header": False,
                    "children": []
                }
            }

            for row_data in rows:
                table_block["table"]["children"].append({
                    "object": "block",
                    "type": "table_row",
                    "table_row": {
                        "cells": row_data
                    }
                })

            blocks.append(table_block)

    return blocks


@app.command("import")
@command
def page_import(
    file_path: str = typer.Argument(
        ...,
        help="Path to the Word document (.docx) to import",
    ),
    parent_page_id: str = typer.Option(
        ...,
        "--parent",
        "-p",
        help="Parent page ID to create the new page under",
    ),
    title: Optional[str] = typer.Option(
        None,
        "--title",
        "-t",
        help="Page title (defaults to filename without extension)",
    ),
    icon: Optional[str] = typer.Option(
        None,
        "--icon",
        help=(
            "Page icon: 'emoji:<shortcode>' (e.g. 'emoji:rocket'), "
            "a literal emoji character (e.g. '🚀'), or 'url:https://...'"
        ),
    ),
):
    """
    Import a Word document (.docx) as a new Notion page.

    Converts Word document content to Notion blocks:
    - Headings → Heading blocks
    - Paragraphs → Paragraph blocks with formatting (bold, italic, etc.)
    - Bullet lists → Bulleted list items
    - Numbered lists → Numbered list items
    - Tables → Table blocks

    Examples:
        notion pages import document.docx --parent PARENT_PAGE_ID
        notion pages import report.docx -p PARENT_PAGE_ID --title "Q4 Report"
        notion pages import notes.docx -p PARENT_PAGE_ID --icon "emoji:page_facing_up"
    """
    try:
        # Validate file exists and is a .docx
        path = Path(file_path)
        if not path.exists():
            print_warning(f"File not found: {file_path}")
            raise typer.Exit(1)

        if path.suffix.lower() != ".docx":
            print_warning(f"File must be a .docx file, got: {path.suffix}")
            raise typer.Exit(1)

        # Determine title
        page_title = title or path.stem

        # Parse icon
        icon_obj = None
        if icon:
            try:
                icon_obj = parse_icon(icon)
            except IconFormatError as exc:
                print_error(str(exc))
                raise typer.Exit(1)

        # Convert Word doc to Notion blocks
        typer.echo(f"Parsing {path.name}...", err=True)
        blocks = docx_to_notion_blocks(str(path))

        if not blocks:
            print_warning("No content found in document")
            raise typer.Exit(1)

        typer.echo(f"Converted {len(blocks)} block(s)", err=True)

        # Create the page
        client = get_client()

        created_page = client.create_standalone_page(
            parent_page_id=parent_page_id,
            title=page_title,
            children=blocks[:100] if len(blocks) <= 100 else None,  # Notion limit
            icon=icon_obj,
        )

        # If more than 100 blocks, append the rest with nesting-aware upload
        if len(blocks) > 100:
            page_id = created_page["id"]
            remaining_blocks = blocks[100:]

            def nesting_progress_cb(stage, message):
                typer.echo(f"  [{stage}] {message}", err=True)

            typer.echo(f"Uploading remaining {len(remaining_blocks)} blocks...", err=True)
            client._upload_blocks_with_nesting(
                page_id, remaining_blocks, progress_callback=nesting_progress_cb
            )

        formatted = format_page_for_display(created_page)
        print_json(formatted)
        print_success(f"Imported '{page_title}' with {len(blocks)} blocks: {created_page.get('url', created_page.get('id'))}")

    except ImportError:
        print_warning("python-docx is required for Word import. Install with: pip install python-docx")
        raise typer.Exit(1)


def markdown_to_styled_html(markdown_content: str, title: str = "Notion Export") -> str:
    """
    Convert markdown content to styled HTML suitable for PDF export.

    Args:
        markdown_content: Markdown text
        title: Document title

    Returns:
        Complete HTML document string
    """
    import re

    # Convert markdown to HTML
    html_body = markdown_content

    # Headings
    html_body = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_body, flags=re.MULTILINE)
    html_body = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html_body, flags=re.MULTILINE)
    html_body = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_body, flags=re.MULTILINE)

    # Bold and italic
    html_body = re.sub(r'\*\*\*(.+?)\*\*\*', r'<strong><em>\1</em></strong>', html_body)
    html_body = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html_body)
    html_body = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html_body)

    # Horizontal rules
    html_body = re.sub(r'^---+$', r'<hr>', html_body, flags=re.MULTILINE)

    # Process lines - collapse multiple blank lines, minimal spacing
    lines = html_body.split('\n')
    processed_lines = []
    in_list = False
    prev_was_blank = False

    for line in lines:
        stripped = line.strip()

        # Skip excessive blank lines (allow max 1)
        if not stripped:
            if not prev_was_blank:
                prev_was_blank = True
            continue

        prev_was_blank = False

        if stripped.startswith('<h') or stripped.startswith('<hr'):
            if in_list:
                processed_lines.append('</ul>')
                in_list = False
            processed_lines.append(stripped)
        elif stripped.startswith('- '):
            if not in_list:
                processed_lines.append('<ul>')
                in_list = True
            processed_lines.append(f'<li>{stripped[2:]}</li>')
        else:
            if in_list:
                processed_lines.append('</ul>')
                in_list = False
            # Wrap in <p> unless it's a block-level element
            if stripped.startswith('<ul') or stripped.startswith('<ol') or stripped.startswith('<li') or stripped.startswith('</'):
                processed_lines.append(stripped)
            else:
                processed_lines.append(f'<p>{stripped}</p>')

    if in_list:
        processed_lines.append('</ul>')

    html_body = '\n'.join(processed_lines)

    # Build complete HTML document with styling
    html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        @page {{
            size: Letter;
            margin: 0.75in;
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            color: #1a1a1a;
            max-width: 100%;
            margin: 0;
            padding: 0;
        }}
        h1 {{
            font-size: 22pt;
            font-weight: 700;
            margin-top: 0;
            margin-bottom: 12pt;
            color: #000;
        }}
        h2 {{
            font-size: 14pt;
            font-weight: 600;
            margin-top: 16pt;
            margin-bottom: 8pt;
            color: #000;
            page-break-after: avoid;
            break-after: avoid;
        }}
        h3 {{
            font-size: 12pt;
            font-weight: 600;
            margin-top: 12pt;
            margin-bottom: 6pt;
            color: #000;
            page-break-after: avoid;
            break-after: avoid;
        }}
        p {{
            margin: 0 0 8pt 0;
        }}
        p:has(> strong:first-child) {{
            page-break-after: avoid;
            break-after: avoid;
        }}
        hr {{
            border: none;
            border-top: 1px solid #d0d0d0;
            margin: 12pt 0;
        }}
        strong {{
            font-weight: 600;
        }}
        ul, ol {{
            margin: 6pt 0;
            padding-left: 20pt;
        }}
        li {{
            margin-bottom: 2pt;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>'''

    return html


# Every export format the command can produce, and the output-file extensions
# that name each one. --output is required, so its extension is an unambiguous
# statement of intent; honoring it is what keeps `-o page.md` from rendering a
# PDF into a file named `.md`.
EXPORT_FORMATS = ("pdf", "html", "md", "notion-json")

EXPORT_FORMAT_BY_SUFFIX = {
    ".pdf": "pdf",
    ".html": "html",
    ".htm": "html",
    ".md": "md",
    ".markdown": "md",
    ".json": "notion-json",
}


def render_html_to_pdf(html_content: str, output_path: Path) -> None:
    """Render HTML to PDF with headless Google Chrome.

    Launches Playwright's ``chrome`` channel, which drives the Google Chrome
    already installed on the machine -- the same binary every other browser-backed
    CLI in this repo uses. Playwright's own bundled Chromium is deliberately not
    used: it is a separate ~150MB download per tool venv, versioned against that
    venv's exact Playwright release, so it silently rots the moment Playwright is
    upgraded and has to be re-fetched by hand. System Chrome has neither problem.
    """
    from playwright.sync_api import Error as PlaywrightError
    from playwright.sync_api import sync_playwright

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(channel="chrome")
            try:
                page = browser.new_page()
                page.set_content(html_content, wait_until="networkidle")
                page.pdf(
                    path=str(output_path),
                    format="Letter",
                    margin={"top": "1in", "right": "1in", "bottom": "1in", "left": "1in"},
                    print_background=True,
                )
            finally:
                browser.close()
    except PlaywrightError as e:
        raise RuntimeError(
            "PDF export needs Google Chrome, which could not be launched: "
            f"{e}. Install Google Chrome, or export a browser-free format "
            "instead (--format md, html, or notion-json)."
        ) from e


def resolve_export_format(format: Optional[str], output_path: Path) -> str:
    """Resolve the export format from --format, or from the output extension.

    An explicit --format always wins. With no --format, the extension of the
    required --output path selects the format. An unrecognized value on either
    side is an error, never a silent default -- exporting the wrong format is
    worse than refusing to guess.
    """
    if format is not None:
        normalized = format.lower()
        if normalized not in EXPORT_FORMATS:
            raise ValueError(
                f"Unknown format: {format}. Use one of: {', '.join(EXPORT_FORMATS)}."
            )
        return normalized

    suffix = output_path.suffix.lower()
    if suffix not in EXPORT_FORMAT_BY_SUFFIX:
        known = ", ".join(sorted(EXPORT_FORMAT_BY_SUFFIX))
        raise ValueError(
            f"Cannot determine export format from output path '{output_path}'. "
            f"Pass --format ({', '.join(EXPORT_FORMATS)}) or use a known "
            f"extension ({known})."
        )
    return EXPORT_FORMAT_BY_SUFFIX[suffix]


@app.command("export")
@command
def page_export(
    page_id: str = typer.Argument(
        ...,
        help="The page ID to export",
    ),
    output: str = typer.Option(
        ...,
        "--output",
        "-o",
        help="Output file path (e.g., document.pdf). Its extension selects the format unless --format is given.",
    ),
    format: Optional[str] = typer.Option(
        None,
        "--format",
        "-f",
        help="Export format: pdf, html, md, or notion-json. Defaults to the --output file extension.",
    ),
):
    """
    Export a Notion page to PDF, HTML, Markdown, or Notion JSON.

    The format comes from the --output extension unless --format overrides it,
    so '-o page.md' writes Markdown and '-o page.pdf' writes a PDF.

    Markdown, HTML, and Notion JSON are built entirely from the Notion REST API
    and need no browser. Only PDF renders through headless Google Chrome.

    The notion-json format exports raw Notion block structures preserving all
    formatting (colors, callouts, columns, bold annotations). The exported JSON
    can be re-imported with 'blocks append --json-file' or 'content set --json-file'.

    Examples:
        notion pages export PAGE_ID --output document.pdf
        notion pages export PAGE_ID -o content.md
        notion pages export PAGE_ID -o content.html
        notion pages export PAGE_ID -o blocks.json
        notion pages export PAGE_ID -o report --format pdf
    """
    output_path = Path(output)
    format_lower = resolve_export_format(format, output_path)

    client = get_client()

    # Get page metadata
    page = client.get_page(page_id)
    page_data = format_page_for_display(page)
    title = page_data.get("title", "Untitled")

    typer.echo(f"Exporting '{title}' as {format_lower}...", err=True)

    if format_lower == "notion-json":
        # Export as clean Notion JSON blocks
        clean_blocks = client.get_blocks_as_notion_json(page_id)

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(clean_blocks, f, indent=2, ensure_ascii=False)

        print_success(f"Exported {len(clean_blocks)} blocks to {output_path}")
        print_json({
            "page_id": page_id,
            "title": title,
            "format": format_lower,
            "output": str(output_path),
            "block_count": len(clean_blocks),
        })
        return

    # Get page content as markdown for other formats
    blocks = client.get_block_children_all(page_id, recursive=True)
    markdown_content = blocks_to_markdown(blocks)

    # Add title as H1
    full_markdown = f"# {title}\n\n{markdown_content}"

    if format_lower == "md":
        # Export as markdown
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_markdown)
        print_success(f"Exported to {output_path}")

    elif format_lower == "html":
        # Export as HTML
        html_content = markdown_to_styled_html(full_markdown, title)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print_success(f"Exported to {output_path}")

    elif format_lower == "pdf":
        html_content = markdown_to_styled_html(full_markdown, title)
        render_html_to_pdf(html_content, output_path)
        print_success(f"Exported to {output_path}")

    print_json({
        "page_id": page_id,
        "title": title,
        "format": format_lower,
        "output": str(output_path),
    })


@app.command("duplicate")
@command
def page_duplicate(
    page_id: str = typer.Argument(
        ...,
        help="The page ID to duplicate",
    ),
    title: Optional[str] = typer.Option(
        None,
        "--title",
        "-t",
        help="New page title (defaults to 'Copy of {original}')",
    ),
    properties: Optional[str] = typer.Option(
        None,
        "--properties",
        "-p",
        help="Raw JSON property overrides (Notion API format)",
    ),
    to_database: Optional[str] = typer.Option(
        None,
        "--to-database",
        help="Target database ID (defaults to same database as source)",
    ),
    replace: Optional[List[str]] = typer.Option(
        None,
        "--replace",
        "-r",
        help="Find/replace in block text (format: 'old:new'). Can be used multiple times.",
    ),
):
    """
    Duplicate a page including all blocks, child pages/subpages, and rich formatting.

    Supports pages in databases and standalone pages (page parents).
    Preserves column layouts, callouts with colors, bold/italic annotations,
    images, and recursively nested child-page content. Handles arbitrary nesting depth.
    Fails before creating the destination when the source contains a block type
    that the Notion API cannot recreate completely.

    Examples:
        notion pages duplicate PAGE_ID
        notion pages duplicate PAGE_ID --title "New Copy"
        notion pages duplicate PAGE_ID --title "2026 Contract" --replace "2025:2026" --replace "$99,400:$84,000"
        notion pages duplicate PAGE_ID --to-database TARGET_DB_ID
        notion pages duplicate PAGE_ID --properties '{"Status": {"status": {"name": "Draft"}}}'
    """
    client = get_client()

    # Parse property overrides
    prop_overrides = None
    if properties:
        try:
            prop_overrides = json.loads(properties)
        except json.JSONDecodeError as e:
            print_warning(f"Invalid JSON in --properties: {e}")
            raise typer.Exit(1)

    # Parse replacements
    replacements = None
    if replace:
        replacements = []
        for r in replace:
            if ":" not in r:
                print_warning(f"Invalid --replace format '{r}'. Use 'old:new' format.")
                raise typer.Exit(1)
            old, new = r.split(":", 1)
            replacements.append((old, new))

    # Progress callback
    def progress_cb(stage, message):
        typer.echo(f"  [{stage}] {message}", err=True)

    typer.echo("Duplicating page...", err=True)

    new_page = client.duplicate_page(
        page_id=page_id,
        title=title,
        property_overrides=prop_overrides,
        target_database_id=to_database,
        replacements=replacements,
        progress_callback=progress_cb,
    )

    formatted = format_page_for_display(new_page)
    print_json(formatted)
    print_success(f"Page duplicated: {new_page.get('url', new_page.get('id'))}")


@app.command("update")
@command
def page_update(
    page_id: str = typer.Argument(
        ...,
        help="The page ID to update",
    ),
    title: Optional[str] = typer.Option(
        None,
        "--title",
        "-t",
        help="New page title",
    ),
    icon: Optional[str] = typer.Option(
        None,
        "--icon",
        help=(
            "Page icon: 'emoji:<shortcode>' (e.g. 'emoji:rocket'), "
            "a literal emoji character (e.g. '🚀'), or 'url:https://...'"
        ),
    ),
    archive: Optional[bool] = typer.Option(
        None,
        "--archive/--restore",
        help="Archive or restore the page",
    ),
):
    """
    Update a page's title, icon, or archive status.

    Examples:
        notion pages update PAGE_ID --title "New Title"
        notion pages update PAGE_ID --icon "emoji:star"
        notion pages update PAGE_ID --icon "⭐"
        notion pages update PAGE_ID --archive
        notion pages update PAGE_ID --restore
    """
    client = get_client()

    properties = None
    icon_obj = None

    # Build title property if provided
    if title:
        properties = {
            "title": {
                "title": [{"type": "text", "text": {"content": title}}]
            }
        }

    # Parse icon
    if icon:
        try:
            icon_obj = parse_icon(icon)
        except IconFormatError as exc:
            print_error(str(exc))
            raise typer.Exit(1)

    # Validate we have something to update
    if properties is None and icon_obj is None and archive is None:
        print_warning("No updates specified. Use --title, --icon, or --archive/--restore")
        raise typer.Exit(1)

    # Perform update
    updated_page = client.update_page(
        page_id=page_id,
        properties=properties,
        archived=archive,
        icon=icon_obj,
    )

    formatted = format_page_for_display(updated_page)
    print_json(formatted)
    print_success(f"Page {page_id} updated successfully.")


@app.command("delete")
@command
def page_delete(
    page_id: str = typer.Argument(
        ...,
        help="The page ID to delete (archive)",
    ),
    force: bool = typer.Option(
        False,
        "--force",
        "-F",
        help="Skip confirmation prompt",
    ),
):
    """
    Delete (archive) a page.

    Note: Notion doesn't truly delete pages - this archives them.
    Archived pages can be restored using 'notion pages update PAGE_ID --restore'.

    Examples:
        notion pages delete abc123-def456
        notion pages delete abc123-def456 --force
    """
    client = get_client()

    # Confirm unless force flag is set
    if not force:
        confirm = typer.confirm(f"Archive page {page_id}?")
        if not confirm:
            typer.echo("Cancelled.")
            raise typer.Exit(0)

    # Archive the page
    updated_page = client.update_page(
        page_id=page_id,
        archived=True,
    )

    formatted = format_page_for_display(updated_page)
    print_json(formatted)
    print_success(f"Page {page_id} archived successfully.")


@content_app.command("append")
@command
def content_append(
    page_id: str = typer.Argument(
        ...,
        help="The page ID to append content to",
    ),
    text: Optional[str] = typer.Option(
        None,
        "--text",
        "-t",
        help="Text/markdown content to append",
    ),
    file: Optional[str] = typer.Option(
        None,
        "--file",
        "-f",
        help="File containing text/markdown content to append",
    ),
    paragraph: Optional[str] = typer.Option(
        None,
        "--paragraph",
        "-p",
        help="Add a simple paragraph block",
    ),
    is_toggleable: bool = typer.Option(
        False,
        "--is-toggleable",
        help="Make all heading_1/2/3 blocks generated from markdown into toggle headings",
    ),
):
    """
    Append content blocks to a page.

    Automatically handles content larger than Notion's 100-block limit
    by splitting into chunks and uploading sequentially.

    Examples:
        notion pages content append PAGE_ID --text "Hello world"
        notion pages content append PAGE_ID --text "# Heading\\n\\nParagraph text"
        notion pages content append PAGE_ID --file content.md
        notion pages content append PAGE_ID --paragraph "Simple paragraph"
        notion pages content append PAGE_ID --file chapters.md --is-toggleable
    """
    try:
        client = get_client()

        # Determine content source
        content = None
        if text:
            content = text
        elif file:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
        elif paragraph:
            content = paragraph

        if not content:
            print_warning("No content specified. Use --text, --file, or --paragraph")
            raise typer.Exit(1)

        # Upload local markdown images, then convert to blocks with the
        # resulting file_upload IDs so ![alt](path) becomes an image block.
        image_uploads = process_markdown_images(content, file, client)
        blocks = text_to_blocks(
            content, image_uploads=image_uploads, is_toggleable=is_toggleable
        )

        if not blocks:
            print_warning("No valid content blocks generated from input")
            raise typer.Exit(1)

        # Define progress callback for nesting-aware upload
        def nesting_progress_cb(stage, message):
            typer.echo(f"  [{stage}] {message}", err=True)

        # Show info if content is large
        if len(blocks) > 100:
            typer.echo(f"Content has {len(blocks)} blocks, uploading...", err=True)

        # Append blocks using nesting-aware method (handles Notion's 2-level nesting limit)
        created_count, _ = client._upload_blocks_with_nesting(
            page_id, blocks, progress_callback=nesting_progress_cb
        )

        print_success(f"Appended {created_count} block(s) to page {page_id}")
        print_json({"blocks_created": created_count, "page_id": page_id})

    except FileNotFoundError:
        print_warning(f"File not found: {file}")
        raise typer.Exit(1)


@content_app.command("set")
@command
def content_set(
    page_id: str = typer.Argument(
        ...,
        help="The page ID to set content for",
    ),
    text: Optional[str] = typer.Option(
        None,
        "--text",
        "-t",
        help="Text/markdown content to set (replaces existing)",
    ),
    file: Optional[str] = typer.Option(
        None,
        "--file",
        "-f",
        help="File containing text/markdown content to set",
    ),
    json_file: Optional[str] = typer.Option(
        None,
        "--json-file",
        help="File containing Notion JSON blocks to set (from 'export --format notion-json')",
    ),
    is_toggleable: bool = typer.Option(
        False,
        "--is-toggleable",
        help="Make all heading_1/2/3 blocks generated from markdown into toggle headings (markdown input only; ignored with --json-file)",
    ),
):
    """
    Replace all page content with new content.

    This clears existing content and sets new content.
    Automatically handles content larger than Notion's 100-block limit
    by splitting into chunks and uploading sequentially.

    Use --json-file to import raw Notion blocks (preserving rich formatting like
    callouts, colors, columns) exported via 'notion pages export --format notion-json'.

    Examples:
        notion pages content set PAGE_ID --text "New content"
        notion pages content set PAGE_ID --file content.md
        notion pages content set PAGE_ID --json-file blocks.json
        notion pages content set PAGE_ID --file outline.md --is-toggleable
    """
    try:
        client = get_client()

        # Validate mutually exclusive options
        sources = sum(1 for x in [text, file, json_file] if x is not None)
        if sources == 0:
            print_warning("No content specified. Use --text, --file, or --json-file")
            raise typer.Exit(1)
        if sources > 1:
            print_warning("Only one of --text, --file, or --json-file can be specified")
            raise typer.Exit(1)

        # Determine content source
        if json_file:
            with open(json_file, 'r', encoding='utf-8') as f:
                blocks = json.load(f)
            if not isinstance(blocks, list):
                print_warning("JSON file must contain an array of block objects")
                raise typer.Exit(1)
        else:
            content = None
            if text:
                content = text
            elif file:
                with open(file, 'r', encoding='utf-8') as f:
                    content = f.read()

            # Upload local markdown images FIRST. This runs before
            # clear_page_content below, so a missing or unsupported image file
            # fails loudly with the page's existing content still intact.
            image_uploads = process_markdown_images(content, file, client)
            blocks = text_to_blocks(
                content, image_uploads=image_uploads, is_toggleable=is_toggleable
            )

        if not blocks:
            print_warning("No valid content blocks generated from input")
            raise typer.Exit(1)

        # Transform and validate the ENTIRE payload BEFORE the destructive clear.
        #
        # Notion rejects any rich_text text.content over 2000 chars and any
        # rich_text array over 100 elements. enforce_block_limits reshapes the
        # payload so every block is uploadable (splitting oversize paragraphs,
        # headings, list items, code, etc. on word boundaries and overflowing
        # into sibling blocks when needed). Doing this first means a pre-checkable
        # validation failure can never empty the page: if the payload cannot be
        # made valid, we exit before touching the page's existing content.
        blocks = enforce_block_limits(blocks)
        invalid = _find_oversize_rich_text(blocks)
        if invalid is not None:
            print_warning(
                "Refusing to clear page: a block exceeds Notion's 2000-character "
                f"rich_text limit and could not be split ({invalid}). "
                "The page's existing content was left untouched."
            )
            raise typer.Exit(1)

        # Clear existing content (single API call using erase_content flag).
        # Only reached once the new payload is known to be uploadable.
        client.clear_page_content(page_id)

        # Upload blocks using nesting-aware method (handles Notion's 2-level nesting limit)
        def nesting_progress_cb(stage, message):
            typer.echo(f"  [{stage}] {message}", err=True)

        if len(blocks) > 100:
            typer.echo(f"Uploading {len(blocks)} blocks...", err=True)

        try:
            created_count, _ = client._upload_blocks_with_nesting(
                page_id, blocks, progress_callback=nesting_progress_cb
            )
        except Exception as upload_error:
            # The page was already cleared. A failure here (e.g. a network drop
            # mid-upload) can leave the page partially populated or empty. Fail
            # loudly with the exact recovery path instead of exiting as if
            # nothing happened.
            print_warning(
                f"Upload failed AFTER the page was cleared: {upload_error}. "
                f"Page {page_id} may now be empty or partially populated. "
                "Re-run 'notion pages content set' with the same input to restore it."
            )
            raise typer.Exit(1)

        print_success(f"Replaced content with {created_count} block(s)")
        print_json({
            "page_id": page_id,
            "blocks_created": created_count,
        })

    except FileNotFoundError as e:
        print_warning(f"File not found: {e.filename or json_file or file}")
        raise typer.Exit(1)
    except json.JSONDecodeError as e:
        print_warning(f"Invalid JSON: {e}")
        raise typer.Exit(1)


_HEADING_LEVELS = {"heading_1": 1, "heading_2": 2, "heading_3": 3}


def _parse_markdown_heading(value: str, label: str) -> tuple:
    """
    Parse '## Title' into (level, text, block_type).

    Exits 1 when the markdown prefix is missing -- the prefix is the only thing
    that identifies which heading level to match.
    """
    stripped = value.strip()
    for prefix, level in (("### ", 3), ("## ", 2), ("# ", 1)):
        if stripped.startswith(prefix):
            return level, stripped[len(prefix):].strip(), f"heading_{level}"

    print_warning(
        f"{label} must start with a markdown prefix (# ## ###). "
        f"Got: '{stripped}'"
    )
    raise typer.Exit(1)


def _heading_block_text(block: dict) -> str:
    """Plain text of a heading block, matching how the markdown heading reads."""
    block_type = block.get("type", "")
    rich_text = block.get(block_type, {}).get("rich_text", [])
    return "".join(rt.get("plain_text", "") for rt in rich_text).strip()


def _section_end_index(all_blocks: list, start_idx: int, level: int) -> int:
    """Index one past the last block of the section starting at start_idx."""
    for i in range(start_idx + 1, len(all_blocks)):
        block_level = _HEADING_LEVELS.get(all_blocks[i].get("type", ""))
        if block_level is not None and block_level <= level:
            return i
    return len(all_blocks)


def _enclosing_heading(all_blocks: list, start_idx: int, level: int) -> Optional[str]:
    """Nearest preceding heading at a higher level, rendered as markdown."""
    for i in range(start_idx - 1, -1, -1):
        block_level = _HEADING_LEVELS.get(all_blocks[i].get("type", ""))
        if block_level is not None and block_level < level:
            return f"{'#' * block_level} {_heading_block_text(all_blocks[i])}"
    return None


def _find_heading_matches(all_blocks: list, block_type: str, text: str, level: int) -> list:
    """
    Every top-level heading block whose type and text match, in page order.

    Each match carries the resolved section range and the enclosing higher-level
    heading so an ambiguous call can report exactly which sections it could mean.
    """
    matches = []
    for i, block in enumerate(all_blocks):
        if block.get("type") != block_type:
            continue
        if _heading_block_text(block) != text:
            continue
        matches.append({
            "occurrence": len(matches) + 1,
            "start_index": i,
            "end_index": _section_end_index(all_blocks, i, level),
            "under": _enclosing_heading(all_blocks, i, level),
        })
    return matches


def _describe_matches(matches: list, total_blocks: int) -> str:
    """Human-readable list of candidate sections for an ambiguity failure."""
    lines = []
    for match in matches:
        under = match["under"] or "(no enclosing heading)"
        lines.append(
            f"  --occurrence {match['occurrence']}: blocks "
            f"{match['start_index'] + 1}-{match['end_index']} of {total_blocks}, "
            f"under {under}"
        )
    return "\n".join(lines)


@content_app.command("replace-section")
@command
def content_replace_section(
    page_id: str = typer.Argument(
        ...,
        help="The page ID containing the section to replace",
    ),
    heading: str = typer.Option(
        ...,
        "--heading",
        "-h",
        help="Section heading to find (e.g., '## My Section Title'). Include the markdown heading prefix (# ## ###).",
    ),
    occurrence: Optional[int] = typer.Option(
        None,
        "--occurrence",
        "-n",
        help="Which match to target when the heading appears more than once (1-based, in page order).",
    ),
    under: Optional[str] = typer.Option(
        None,
        "--under",
        "-u",
        help="Restrict the search to the section under this higher-level heading (e.g., '## Phase 4'). Include the markdown prefix.",
    ),
    text: Optional[str] = typer.Option(
        None,
        "--text",
        "-t",
        help="New markdown content to replace the section with",
    ),
    file: Optional[str] = typer.Option(
        None,
        "--file",
        "-f",
        help="File containing new markdown content to replace the section with",
    ),
    dry_run: bool = typer.Option(
        False,
        "--dry-run",
        help="Show what would be replaced without making changes",
    ),
):
    """
    Replace a specific section of a page without touching the rest.

    Finds the section by its heading text, then replaces all blocks from that
    heading to (but not including) the next heading at the same or higher level.

    The heading argument must include the markdown prefix to identify the heading
    level: '# Title' for H1, '## Title' for H2, '### Title' for H3.

    When the heading text appears more than once on the page, the command FAILS
    and lists every match rather than guessing. Disambiguate with --occurrence
    (1-based, in page order) or --under (scope the search to one parent section).

    Examples:
        notion pages content replace-section PAGE_ID --heading "## My Section" --file updated.md
        notion pages content replace-section PAGE_ID --heading "### Subsection" --text "### Subsection\\n\\nNew content here"
        notion pages content replace-section PAGE_ID --heading "## Old Title" --file new.md --dry-run
        notion pages content replace-section PAGE_ID --heading "### Your Actions" --occurrence 4 --file new.md
        notion pages content replace-section PAGE_ID --heading "### Your Actions" --under "## Phase 4" --file new.md
    """
    import concurrent.futures

    try:
        client = get_client()

        # Validate content source
        if not text and not file:
            print_warning("No content specified. Use --text or --file")
            raise typer.Exit(1)
        if text and file:
            print_warning("Only one of --text or --file can be specified")
            raise typer.Exit(1)

        # Parse the heading to determine level and text
        heading_stripped = heading.strip()
        target_level, target_text, target_block_type = _parse_markdown_heading(
            heading_stripped, "Heading"
        )

        # --under scopes the search to one parent section, so it must name a
        # heading ABOVE the target level. An equal or lower level cannot contain
        # the target section.
        under_stripped = None
        if under is not None:
            under_stripped = under.strip()
            under_level, under_text, under_block_type = _parse_markdown_heading(
                under_stripped, "Parent heading (--under)"
            )
            if under_level >= target_level:
                print_warning(
                    f"--under must name a higher-level heading than --heading. "
                    f"'{under_stripped}' is H{under_level} and "
                    f"'{heading_stripped}' is H{target_level}."
                )
                raise typer.Exit(1)

        # Load new content
        content = None
        if text:
            content = text
        elif file:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()

        # --dry-run must stay read-only, so no upload happens on that path.
        image_uploads = {}
        if not dry_run:
            image_uploads = process_markdown_images(content, file, client)

        new_blocks = text_to_blocks(content, image_uploads=image_uploads)
        if not new_blocks:
            print_warning("No valid content blocks generated from input")
            raise typer.Exit(1)

        # Transform and validate the ENTIRE new payload BEFORE any mutation.
        #
        # replace-section is NOT atomic: it inserts the new blocks and then
        # deletes the old ones. If the Notion API rejected a new block mid-insert
        # (e.g. an unsupported code-fence language such as ```kql, or a rich_text
        # value over the 2000-char limit), the section was left half-written AND
        # the old section was never deleted -- producing a duplicate heading and
        # a partial section. enforce_block_limits reshapes oversize content and
        # normalizes unsupported code languages to "plain text"; the validator
        # then confirms nothing is left that the API would reject. Doing this
        # before touching the page means a pre-checkable failure can never
        # corrupt the section: if the payload cannot be made valid, we exit
        # before deleting or inserting anything.
        new_blocks = enforce_block_limits(new_blocks)
        invalid = _find_oversize_rich_text(new_blocks)
        if invalid is not None:
            print_warning(
                "Refusing to modify section: a block exceeds Notion's "
                f"2000-character rich_text limit and could not be split ({invalid}). "
                "The page's existing content was left untouched."
            )
            raise typer.Exit(1)

        # Fetch all top-level child blocks of the page (non-recursive, we only need top-level)
        typer.echo("Fetching page blocks...", err=True)
        all_blocks = client.get_block_children_all(page_id, recursive=False)

        if not all_blocks:
            print_warning("Page has no content blocks")
            raise typer.Exit(1)

        # Find every heading block that matches. A page can repeat the same
        # heading text under several parents (e.g. "### Your Actions" under four
        # phases). Silently taking the first match would rewrite the wrong
        # section and destroy content, so ambiguity is a hard failure.
        matches = _find_heading_matches(
            all_blocks, target_block_type, target_text, target_level
        )

        if under_stripped is not None:
            under_matches = _find_heading_matches(
                all_blocks, under_block_type, under_text, under_level
            )
            if not under_matches:
                print_warning(
                    f"Parent heading not found: '{under_stripped}'\n"
                    "Make sure the heading text matches exactly (case-sensitive)."
                )
                raise typer.Exit(1)
            if len(under_matches) > 1:
                print_warning(
                    f"Parent heading '{under_stripped}' matches "
                    f"{len(under_matches)} sections on page {page_id}, so --under "
                    "cannot identify one scope:\n"
                    f"{_describe_matches(under_matches, len(all_blocks))}\n"
                    "Use a --under heading that appears once."
                )
                raise typer.Exit(1)

            scope = under_matches[0]
            matches = [
                match for match in matches
                if scope["start_index"] < match["start_index"] < scope["end_index"]
            ]
            # Renumber so --occurrence counts within the scoped section.
            for position, match in enumerate(matches, start=1):
                match["occurrence"] = position

        if not matches:
            scope_note = f" under '{under_stripped}'" if under_stripped else ""
            print_warning(
                f"Section heading not found{scope_note}: '{heading_stripped}'\n"
                "Make sure the heading text matches exactly (case-sensitive)."
            )
            raise typer.Exit(1)

        if occurrence is not None:
            if occurrence < 1 or occurrence > len(matches):
                print_warning(
                    f"--occurrence {occurrence} is out of range: "
                    f"'{heading_stripped}' matches {len(matches)} section(s) on "
                    f"page {page_id}:\n"
                    f"{_describe_matches(matches, len(all_blocks))}"
                )
                raise typer.Exit(1)
            match = matches[occurrence - 1]
        elif len(matches) > 1:
            print_warning(
                f"'{heading_stripped}' matches {len(matches)} sections on page "
                f"{page_id}. Refusing to guess which one to replace:\n"
                f"{_describe_matches(matches, len(all_blocks))}\n"
                "Re-run with --occurrence N, or --under '<parent heading>' to "
                "scope the search to one parent section."
            )
            raise typer.Exit(1)
        else:
            match = matches[0]

        section_start_idx = match["start_index"]
        section_end_idx = match["end_index"]
        resolved_occurrence = match["occurrence"]

        # Collect block IDs to delete
        section_block_ids = [all_blocks[i]["id"] for i in range(section_start_idx, section_end_idx)]

        # Determine the "after" block ID for insertion.
        #
        # Notion cannot insert before the first child. For a first-child section,
        # insert the replacement after the old section's last block, then delete
        # the old section. The replacement remains before the following blocks
        # without rebuilding the rest of the page.
        if section_start_idx > 0:
            after_block_id = all_blocks[section_start_idx - 1]["id"]
        else:
            after_block_id = all_blocks[section_end_idx - 1]["id"]

        section_block_count = len(section_block_ids)
        blocks_to_delete_count = section_block_count
        typer.echo(
            f"Found section '{target_text}' with {section_block_count} block(s) "
            f"(blocks {section_start_idx + 1}-{section_end_idx} of {len(all_blocks)})",
            err=True,
        )

        if dry_run:
            typer.echo(
                f"[DRY RUN] Would insert {len(new_blocks)} new block(s) and "
                f"delete {blocks_to_delete_count} old block(s)",
                err=True,
            )
            print_json({
                "page_id": page_id,
                "section_heading": heading_stripped,
                "occurrence": resolved_occurrence,
                "total_matches": len(matches),
                "section_start_block": section_start_idx + 1,
                "section_end_block": section_end_idx,
                "section_blocks_to_replace": section_block_count,
                "blocks_to_delete": blocks_to_delete_count,
                "blocks_to_insert": len(new_blocks),
                "first_child_rewrite": False,
                "following_blocks_to_recreate": 0,
                "after_block_id": after_block_id,
                "dry_run": True,
            })
            return

        block_ids_to_delete = list(section_block_ids)

        # Step 1: Insert replacement while the anchor block still exists.
        #
        # The payload was already normalized + validated above, so the API will
        # not reject these blocks for block-limit or code-language reasons. A
        # failure here can only come from a transport-level fault (network drop,
        # timeout). Surface it loudly with the exact recovery path instead of
        # leaving a silently duplicated section: the old section still exists
        # (it is deleted in Step 2), so a partial insert here would create the
        # duplicate-heading state. Re-running with the same input restores it.
        typer.echo(f"Inserting {len(new_blocks)} new block(s)...", err=True)

        def nesting_progress_cb(stage, message):
            typer.echo(f"  [{stage}] {message}", err=True)

        try:
            created_count, _ = client._upload_blocks_with_nesting(
                page_id, new_blocks, progress_callback=nesting_progress_cb, after=after_block_id
            )
        except Exception as insert_error:
            print_warning(
                f"Insert failed while replacing section '{target_text}': {insert_error}. "
                f"The old section was NOT deleted, so page {page_id} may now show a "
                "partial duplicate of this section. Re-run the same "
                "'notion pages content replace-section' command to restore it."
            )
            raise typer.Exit(1)

        # Step 2: Delete old section blocks (parallel for speed).
        #
        # Use the idempotent delete: archiving a top-level block auto-archives
        # its descendants, so a parallel batch that also targets one of those
        # descendants would otherwise get a 400 "Can't edit block that is
        # archived" AFTER the section has already been swapped -- making a
        # successful replace exit non-zero (and skip the caller's verify).
        # delete_block_if_present treats that already-archived case as
        # already-gone (the desired end state) while still raising on every
        # other error, so genuine failures still fail fast.
        typer.echo(f"Deleting {len(block_ids_to_delete)} old block(s)...", err=True)

        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            futures = {
                executor.submit(client.delete_block_if_present, bid): bid
                for bid in block_ids_to_delete
            }
            for future in concurrent.futures.as_completed(futures):
                future.result()  # Raise on any non-already-archived failure

        print_success(f"Replaced section '{target_text}' with {created_count} new block(s)")
        print_json({
            "page_id": page_id,
            "section_heading": heading_stripped,
            "occurrence": resolved_occurrence,
            "total_matches": len(matches),
            "section_start_block": section_start_idx + 1,
            "section_end_block": section_end_idx,
            "section_blocks_replaced": section_block_count,
            "blocks_deleted": len(block_ids_to_delete),
            "blocks_created": created_count,
            "first_child_rewrite": False,
            "following_blocks_recreated": 0,
        })

    except FileNotFoundError:
        print_warning(f"File not found: {file}")
        raise typer.Exit(1)


@content_app.command("clear")
@command
def content_clear(
    page_id: str = typer.Argument(
        ...,
        help="The page ID to clear content from",
    ),
    force: bool = typer.Option(
        False,
        "--force",
        "-F",
        help="Skip confirmation prompt",
    ),
):
    """
    Clear all content from a page.

    Examples:
        notion pages content clear PAGE_ID
        notion pages content clear PAGE_ID --force
    """
    client = get_client()

    # Confirm unless force flag is set
    if not force:
        confirm = typer.confirm(f"Clear all content from page {page_id}?")
        if not confirm:
            typer.echo("Cancelled.")
            raise typer.Exit(0)

    # Clear content (single API call using erase_content flag)
    client.clear_page_content(page_id)

    print_success(f"Cleared all content from page {page_id}")
    print_json({"page_id": page_id, "cleared": True})


# =============================================================================
# Block Commands - Manage individual blocks
# =============================================================================


def _nest_section_under_heading(client, heading_block_id: str, progress_callback=None):
    """
    Move a heading's "section" siblings to be children of the heading.

    A heading's section = blocks that come AFTER the heading in its parent
    container, up to (but not including) the next heading at the same or
    higher level. After making a heading toggleable, those siblings are the
    blocks readers expect to see when they expand the toggle — but Notion's
    API treats them as parent-level siblings, so the toggle would render
    empty unless we re-parent them.

    Notion has no "move block" operation; we re-create the section blocks
    as children of the heading and delete the originals.

    Args:
        client: NotionClient instance.
        heading_block_id: The heading block ID. Must be heading_1/2/3.
        progress_callback: Optional callable(stage, message) for progress.

    Returns:
        Tuple (created_count, deleted_count). Both 0 when there is no
        section content to move (heading at end of parent, or only followed
        by same/higher-level headings).
    """
    import concurrent.futures

    heading_block = client.get_block(heading_block_id)
    heading_type = heading_block.get("type", "")

    if heading_type not in _HEADING_LEVELS:
        return (0, 0)

    target_level = _HEADING_LEVELS[heading_type]

    parent = heading_block.get("parent", {}) or {}
    parent_type = parent.get("type", "")
    if parent_type == "page_id":
        parent_id = parent.get("page_id")
    elif parent_type == "block_id":
        parent_id = parent.get("block_id")
    else:
        return (0, 0)

    if not parent_id:
        return (0, 0)

    siblings = client.get_block_children_all(parent_id, recursive=False)
    if not siblings:
        return (0, 0)

    heading_idx = None
    for i, sibling in enumerate(siblings):
        if sibling.get("id") == heading_block_id:
            heading_idx = i
            break

    if heading_idx is None:
        return (0, 0)

    section_end_idx = len(siblings)
    for i in range(heading_idx + 1, len(siblings)):
        sibling_type = siblings[i].get("type", "")
        sibling_level = _HEADING_LEVELS.get(sibling_type)
        if sibling_level is not None and sibling_level <= target_level:
            section_end_idx = i
            break

    section_blocks = siblings[heading_idx + 1:section_end_idx]
    if not section_blocks:
        return (0, 0)

    if progress_callback:
        progress_callback(
            "nest",
            f"Fetching {len(section_blocks)} section block(s) with nested children...",
        )

    # Hydrate any nested children so we re-create the full subtree.
    client._fetch_children_parallel(section_blocks, max_workers=5)

    # Re-upload Notion-hosted files (image/video/pdf/file) via the File Upload
    # API BEFORE cleaning. Cleaning rewrites a hosted file block to
    # image.external using its signed, EXPIRING S3 URL, which Notion rejects on
    # create (400 image.file_upload should be defined at depth-2) or silently
    # empties to ![]() (depth-1). Re-uploading converts each hosted file to a
    # file_upload reference so the re-parented section keeps its images intact,
    # mirroring the proven `pages duplicate` path.
    reuploaded = client._reupload_file_blocks(section_blocks, progress_callback)
    if reuploaded > 0 and progress_callback:
        progress_callback(
            "nest",
            f"Re-uploaded {reuploaded} Notion-hosted file(s) so they survive re-parenting...",
        )

    cleaned_blocks = client._clean_blocks_recursive(section_blocks)
    if not cleaned_blocks:
        # Section consisted entirely of uncreatable blocks. Don't delete the
        # originals — that would lose data — just bail.
        return (0, 0)

    if progress_callback:
        progress_callback(
            "nest",
            f"Appending {len(cleaned_blocks)} block(s) as children of heading...",
        )

    created_count, _ = client._upload_blocks_with_nesting(
        heading_block_id, cleaned_blocks, progress_callback=progress_callback,
    )

    section_block_ids = [b["id"] for b in section_blocks]
    if progress_callback:
        progress_callback(
            "nest",
            f"Deleting {len(section_block_ids)} original sibling block(s)...",
        )

    # Idempotent delete: these top-level siblings can have children that Notion
    # auto-archives when their parent is archived, so a parallel batch that also
    # targets a descendant must treat its already-archived 400 as already-gone
    # rather than aborting after the re-parent already succeeded. Every other
    # error still raises (fail-fast).
    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        futures = {
            executor.submit(client.delete_block_if_present, bid): bid
            for bid in section_block_ids
        }
        for future in concurrent.futures.as_completed(futures):
            future.result()

    return (created_count, len(section_block_ids))


@blocks_app.command("list")
@command
def blocks_list(
    page_id: str = typer.Option(
        ...,
        "--page-id",
        "-p",
        help="The page or block ID to list children of",
    ),
    recursive: bool = typer.Option(
        False,
        "--recursive",
        "-r",
        help="Recursively fetch all nested blocks",
    ),
    table: bool = typer.Option(
        False,
        "--table",
        "-t",
        help="Display as formatted table",
    ),
    markdown: bool = typer.Option(
        False,
        "--markdown",
        "-m",
        help="Output blocks as markdown",
    ),
    limit: Optional[int] = typer.Option(
        None,
        "--limit",
        "-l",
        help="Maximum number of blocks to return (default: all blocks)",
    ),
    filter: Optional[List[str]] = typer.Option(
        None,
        "--filter",
        "-f",
        help="Filter: field:op:value (e.g., type:eq:paragraph, has_children:eq:true)",
    ),
    properties: Optional[str] = typer.Option(
        None,
        "--properties",
        help="Comma-separated list of fields to include in output",
    ),
):
    """
    List blocks (children) of a page or block.

    Returns the complete child-block list by default, paginating through every
    page of the Notion children endpoint (has_more/next_cursor). Pass --limit to
    cap the result; without it, all blocks are returned.

    Examples:
        notion pages blocks list --page-id PAGE_ID
        notion pages blocks list --page-id PAGE_ID --table
        notion pages blocks list --page-id PAGE_ID --recursive
        notion pages blocks list --page-id PAGE_ID --markdown
        notion pages blocks list --page-id PAGE_ID --limit 50
    """
    client = get_client()
    # Default (limit=None) fetches the COMPLETE child list via full
    # has_more/next_cursor pagination. Only an explicit --limit caps the
    # fetch, and recursive runs always fetch every top-level block.
    fetch_limit = None if recursive else limit
    blocks = client.get_block_children_all(page_id, recursive=recursive, limit=fetch_limit)

    # Format for display first
    formatted = [format_block_for_display(b) for b in blocks]

    # Apply client-side filtering
    if filter:
        formatted = apply_filters(formatted, filter)

    # Apply client-side limit only when the caller asked for one
    if limit is not None:
        formatted = formatted[:limit]

    # Filter to requested properties if specified
    if properties:
        props_list = [p.strip() for p in properties.split(",")]
        filtered_blocks = []
        for block in formatted:
            filtered = {prop: block.get(prop) for prop in props_list if prop in block}
            filtered_blocks.append(filtered)
        formatted = filtered_blocks

    if markdown:
        # For markdown, we need the original blocks structure
        markdown_blocks = blocks if limit is None else blocks[:limit]
        markdown_content = blocks_to_markdown(markdown_blocks)
        typer.echo(markdown_content)
        return

    if table:
        columns = ["id", "type", "text", "has_children", "is_toggleable"] if not properties else props_list
        print_table(
            formatted,
            columns=columns,
        )
    else:
        print_json(formatted)

    typer.echo(f"\n{len(formatted)} block(s) found.", err=True)


@blocks_app.command("get")
@command
def blocks_get(
    block_id: str = typer.Argument(
        ...,
        help="The block ID to retrieve",
    ),
    markdown: bool = typer.Option(
        False,
        "--markdown",
        "-m",
        help="Output block as markdown",
    ),
    include_children: bool = typer.Option(
        False,
        "--children",
        "-c",
        help="Include child blocks",
    ),
):
    """
    Get a specific block by ID.

    Examples:
        notion pages blocks get BLOCK_ID
        notion pages blocks get BLOCK_ID --markdown
        notion pages blocks get BLOCK_ID --children
    """
    client = get_client()
    block = client.get_block(block_id)

    if markdown:
        markdown_content = block_to_markdown(block)
        if include_children and block.get("has_children"):
            children = client.get_block_children_all(block_id, recursive=True)
            markdown_content += "\n" + blocks_to_markdown(children)
        typer.echo(markdown_content)
        return

    result = block
    if include_children and block.get("has_children"):
        children = client.get_block_children_all(block_id, recursive=True)
        result["children"] = children

    print_json(result)


@blocks_app.command("update")
@command
def blocks_update(
    block_id: str = typer.Argument(
        ...,
        help="The block ID to update",
    ),
    text: Optional[str] = typer.Option(
        None,
        "--text",
        "-t",
        help="New PLAIN-TEXT content for text-based blocks (paragraph, heading, etc.). "
        "Markdown syntax is stored literally; use --markdown/--md-file for links and inline code.",
    ),
    markdown: Optional[str] = typer.Option(
        None,
        "--markdown",
        "-m",
        help="New MARKDOWN content; parses [text](url) links, `inline code`, **bold**, "
        "and *italic* into real Notion rich_text in place.",
    ),
    md_file: Optional[str] = typer.Option(
        None,
        "--md-file",
        help="File containing MARKDOWN content for the block (parsed like --markdown).",
    ),
    checked: Optional[bool] = typer.Option(
        None,
        "--checked/--unchecked",
        help="Set checked state for to_do blocks",
    ),
    json_data: Optional[str] = typer.Option(
        None,
        "--json",
        "-j",
        help="Raw JSON block data to update (advanced)",
    ),
    json_file: Optional[str] = typer.Option(
        None,
        "--json-file",
        "-f",
        help="File containing raw JSON block data to update (advanced)",
    ),
    toggleable: Optional[bool] = typer.Option(
        None,
        "--toggleable/--no-toggleable",
        help="For heading_1/2/3 blocks: turn the toggle arrow on or off",
    ),
    no_nest: bool = typer.Option(
        False,
        "--no-nest",
        help=(
            "When using --toggleable, skip auto-nesting of the heading's "
            "section siblings as children. Default behavior moves the blocks "
            "between this heading and the next same/higher-level heading "
            "into the toggle so the expanded toggle actually shows them."
        ),
    ),
):
    """
    Update a block's content.

    For simple text updates, use --text. For complex updates, use --json or --json-file.

    Use --toggleable / --no-toggleable to flip the is_toggleable flag on existing
    heading_1/2/3 blocks. The flag can be combined with --text to rename and
    toggle in a single call.

    By default, --toggleable also re-parents the heading's section siblings
    (everything from the heading down to the next same/higher-level heading)
    so they become children of the toggle. Re-parenting recreates those blocks
    via the API, which means new block IDs and dropped block-scoped comments;
    use --no-nest to opt out and only flip the toggle flag.

    Examples:
        notion pages blocks update BLOCK_ID --text "Updated paragraph text"
        notion pages blocks update BLOCK_ID --checked
        notion pages blocks update BLOCK_ID --toggleable
        notion pages blocks update BLOCK_ID --toggleable --no-nest
        notion pages blocks update BLOCK_ID --no-toggleable
        notion pages blocks update BLOCK_ID --text "New title" --toggleable
        notion pages blocks update BLOCK_ID --json '{"paragraph":{"rich_text":[{"text":{"content":"New text"}}]}}'
        notion pages blocks update BLOCK_ID --json-file block_update.json
    """
    try:
        client = get_client()

        # Determine update data
        update_data = None

        content_sources = [
            name
            for name, value in (
                ("--text", text),
                ("--markdown", markdown),
                ("--md-file", md_file),
                ("--json", json_data),
                ("--json-file", json_file),
            )
            if value is not None
        ]
        if len(content_sources) > 1:
            print_warning(
                "Specify only one content source: --text, --markdown, --md-file, --json, or --json-file"
            )
            raise typer.Exit(1)

        if json_file:
            with open(json_file, 'r', encoding='utf-8') as f:
                update_data = json.load(f)
        elif json_data:
            update_data = json.loads(json_data)
        elif md_file:
            with open(md_file, 'r', encoding='utf-8') as f:
                markdown = f.read()
        if markdown is not None or text is not None or checked is not None or toggleable is not None:
            # Get current block to determine type
            current_block = client.get_block(block_id)
            block_type = current_block.get("type")

            if block_type not in TEXT_EDITABLE_BLOCK_TYPES:
                print_warning(f"Block type '{block_type}' does not support simple text updates. Use --json instead.")
                raise typer.Exit(1)

            if toggleable is not None and block_type not in ("heading_1", "heading_2", "heading_3"):
                print_warning(
                    f"--toggleable / --no-toggleable only applies to heading_1/2/3 blocks (got '{block_type}')."
                )
                raise typer.Exit(1)

            # Build the shared text/checked payload (checked applies only to to_do).
            update_data = build_text_block_update(
                block_type,
                text=markdown if markdown is not None else text,
                checked=checked if block_type == "to_do" else None,
                markdown=markdown is not None,
            )

            if toggleable is not None:
                update_data[block_type]["is_toggleable"] = toggleable
                # Notion's PATCH /blocks/{id} for headings requires rich_text in the
                # body even when only flipping is_toggleable. Preserve existing text
                # if no replacement rich_text was also provided.
                if text is None and markdown is None:
                    existing_rich_text = current_block.get(block_type, {}).get("rich_text", [])
                    update_data[block_type]["rich_text"] = existing_rich_text

        if not update_data:
            print_warning(
                "No update data specified. Use --text, --markdown, --md-file, "
                "--checked, --toggleable, --json, or --json-file"
            )
            raise typer.Exit(1)

        # Perform update
        updated_block = client.update_block(block_id, update_data)

        # When toggling a heading ON, re-parent its section siblings as children
        # so the toggle actually contains the content readers expect to see.
        nested_count = 0
        deleted_count = 0
        if toggleable is True and not no_nest:
            updated_type = updated_block.get("type", "")
            if updated_type in ("heading_1", "heading_2", "heading_3"):
                def _nest_progress(stage, message):
                    typer.echo(f"  [{stage}] {message}", err=True)

                nested_count, deleted_count = _nest_section_under_heading(
                    client, block_id, progress_callback=_nest_progress
                )

        print_json(updated_block)
        if nested_count:
            print_success(
                f"Block {block_id} updated successfully. "
                f"Nested {nested_count} block(s) under the toggle "
                f"(replaced {deleted_count} sibling block(s))."
            )
        else:
            print_success(f"Block {block_id} updated successfully.")

    except FileNotFoundError as e:
        print_warning(f"File not found: {e.filename}")
        raise typer.Exit(1)
    except json.JSONDecodeError as e:
        print_warning(f"Invalid JSON: {e}")
        raise typer.Exit(1)


@blocks_app.command("delete")
@command
def blocks_delete(
    block_id: str = typer.Argument(
        ...,
        help="The block ID to delete",
    ),
    recursive: bool = typer.Option(
        False,
        "--recursive",
        "-r",
        help="Delete all children blocks first (in parallel)",
    ),
    force: bool = typer.Option(
        False,
        "--force",
        "-F",
        help="Skip confirmation prompt",
    ),
):
    """
    Delete a block.

    Note: Deleted blocks are moved to trash and can be recovered via Notion UI.

    Examples:
        notion pages blocks delete BLOCK_ID
        notion pages blocks delete BLOCK_ID --force
        notion pages blocks delete BLOCK_ID --recursive --force
    """
    import concurrent.futures

    client = get_client()

    # Get block info
    block = client.get_block(block_id)
    block_type = block.get("type", "unknown")
    has_children = block.get("has_children", False)

    # Count children if recursive
    children_count = 0
    children_ids = []
    if recursive and has_children:
        children = client.get_block_children_all(block_id, recursive=True)
        children_ids = [c["id"] for c in children]
        children_count = len(children_ids)

    target_description = f"{block_type} block {block_id}"
    if recursive and children_count > 0:
        target_description += f" and {children_count} child block(s)"

    confirm_destructive_action(
        f"Delete {target_description}?",
        assume_yes=force,
        action_description=f"delete {target_description}",
        skip_flag_hint="--force",
    )

    deleted_count = 0
    already_gone_count = 0

    # Delete children in parallel if recursive.
    #
    # Use the idempotent delete: archiving a parent block auto-archives its
    # descendants, so a parallel batch that also targets one of those
    # descendants would otherwise get a 400 "Can't edit block that is
    # archived". delete_block_if_present treats that one already-archived
    # case as already-gone (the desired end state) while still raising on
    # EVERY other error (auth, not-found, transport, any non-archived 400).
    # The result iterator is consumed with list(), so any such error
    # propagates out of the executor and is reported by the outer
    # `except Exception as e: handle_error(e)` with a non-zero exit -- a
    # genuine child-delete failure is never silently dropped.
    if recursive and children_ids:
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            results = list(
                executor.map(client.delete_block_if_present, children_ids)
            )

        # results[i] is True if this run archived the child, False if it was
        # already gone (cascade). Count both truthfully -- the total still
        # accounts for every requested child, with no hidden failures.
        deleted_count = sum(1 for archived in results if archived)
        already_gone_count = len(results) - deleted_count

        if already_gone_count:
            typer.echo(
                f"Deleted {deleted_count} child block(s) "
                f"({already_gone_count} already gone).",
                err=True,
            )
        else:
            typer.echo(f"Deleted {deleted_count} child block(s).", err=True)

    # Delete the main block
    deleted_block = client.delete_block(block_id)

    print_json(format_block_for_display(deleted_block))
    if recursive and children_ids:
        print_success(
            f"Block {block_id} and {len(children_ids)} child block(s) deleted "
            f"successfully ({deleted_count} archived this run, "
            f"{already_gone_count} already gone)."
        )
    else:
        print_success(f"Block {block_id} deleted successfully.")


@blocks_app.command("append")
@command
def blocks_append(
    block_id: str = typer.Argument(
        ...,
        help="The block or page ID to append children to",
    ),
    text: Optional[str] = typer.Option(
        None,
        "--text",
        "-t",
        help="Text/markdown content to append as child blocks",
    ),
    file: Optional[str] = typer.Option(
        None,
        "--file",
        "-f",
        help="File containing text/markdown content to append",
    ),
    json_data: Optional[str] = typer.Option(
        None,
        "--json",
        "-j",
        help="Raw JSON array of blocks to append (advanced)",
    ),
    json_file: Optional[str] = typer.Option(
        None,
        "--json-file",
        help="File containing raw JSON array of blocks to append (advanced)",
    ),
    after: Optional[str] = typer.Option(
        None,
        "--after",
        "-a",
        help="Block ID to insert after (for positioning within parent's children)",
    ),
    is_toggleable: bool = typer.Option(
        False,
        "--is-toggleable",
        help="Make all heading_1/2/3 blocks generated from markdown into toggle headings (markdown input only; ignored with --json/--json-file)",
    ),
    synced: bool = typer.Option(
        False,
        "--synced",
        help="Wrap the supplied content in one original synced block",
    ),
    synced_from: Optional[str] = typer.Option(
        None,
        "--synced-from",
        help="Create a duplicate synced block from the original synced block ID",
    ),
):
    """
    Append child blocks to a block or page.

    Automatically handles content larger than Notion's 100-block limit
    by splitting into chunks and uploading sequentially.

    Examples:
        notion pages blocks append BLOCK_ID --text "New paragraph"
        notion pages blocks append BLOCK_ID --file content.md
        notion pages blocks append BLOCK_ID --json '[{"paragraph":{"rich_text":[{"text":{"content":"Hello"}}]}}]'
        notion pages blocks append PAGE_ID --after BLOCK_ID --text "Inserted after specific block"
        notion pages blocks append PAGE_ID --file chapter.md --is-toggleable
        notion pages blocks append PAGE_ID --file synced.md --synced
        notion pages blocks append PAGE_ID --synced-from ORIGINAL_SYNCED_BLOCK_ID
    """
    try:
        # Determine content source
        blocks = None
        content_sources = [
            name for name, value in (
                ("--text", text),
                ("--file", file),
                ("--json", json_data),
                ("--json-file", json_file),
            )
            if value is not None
        ]

        if synced_from:
            if synced or content_sources:
                print_warning("--synced-from cannot be combined with --synced or content inputs")
                raise typer.Exit(1)
            blocks = [{
                "object": "block",
                "type": "synced_block",
                "synced_block": {
                    "synced_from": {
                        "type": "block_id",
                        "block_id": synced_from,
                    },
                },
            }]
        elif synced and len(content_sources) != 1:
            print_warning("--synced requires exactly one content input: --text, --file, --json, or --json-file")
            raise typer.Exit(1)

        # Built before block conversion: markdown input uploads its local images
        # through this client before anything is appended. Option-validation
        # failures above exit without ever creating a client.
        client = get_client()

        if blocks is None and json_file:
            with open(json_file, 'r', encoding='utf-8') as f:
                blocks = json.load(f)
        elif blocks is None and json_data:
            blocks = json.loads(json_data)
        elif blocks is None and text:
            image_uploads = process_markdown_images(text, None, client)
            blocks = text_to_blocks(
                text, image_uploads=image_uploads, is_toggleable=is_toggleable
            )
        elif blocks is None and file:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            image_uploads = process_markdown_images(content, file, client)
            blocks = text_to_blocks(
                content, image_uploads=image_uploads, is_toggleable=is_toggleable
            )

        if not blocks:
            print_warning("No content specified. Use --text, --file, --json, or --json-file")
            raise typer.Exit(1)

        if synced:
            blocks = [{
                "object": "block",
                "type": "synced_block",
                "synced_block": {
                    "synced_from": None,
                    "children": blocks,
                },
            }]

        # Define progress callback for nesting-aware upload
        def nesting_progress_cb(stage, message):
            typer.echo(f"  [{stage}] {message}", err=True)

        # Show info if content is large
        if len(blocks) > 100:
            typer.echo(f"Content has {len(blocks)} blocks, uploading...", err=True)

        # One upload path for both positioned (--after) and plain appends.
        # _upload_blocks_with_nesting chunks past Notion's 100-block-per-call
        # limit, handles arbitrary nesting depth, threads `after` into only the
        # first chunk to position the insertion, and returns the true count of
        # blocks created (top-level + descendants). It never re-fetches or
        # recreates the existing tail, so an --after insert past position 100
        # cannot drop blocks. (The raw single-call append returns Notion's full
        # repositioned tail in `results`, which would both overcount and cap at
        # 100 -- avoided by routing through the chunked, nesting-aware uploader.)
        created_count, _ = client._upload_blocks_with_nesting(
            block_id, blocks, progress_callback=nesting_progress_cb, after=after
        )

        print_success(f"Appended {created_count} block(s) to {block_id}")
        print_json({"blocks_created": created_count, "parent_id": block_id})

    except FileNotFoundError as e:
        print_warning(f"File not found: {e.filename}")
        raise typer.Exit(1)
    except json.JSONDecodeError as e:
        print_warning(f"Invalid JSON: {e}")
        raise typer.Exit(1)
