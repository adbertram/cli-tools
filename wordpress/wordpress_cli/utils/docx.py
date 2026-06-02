"""DOCX file conversion utilities for WordPress.

This module handles conversion of Word documents (.docx) to HTML suitable for
WordPress publication, including:
- Extraction of post title from document headings
- Conversion of .docx content to HTML using mammoth
- Automatic image upload via callback pattern
- Automatic link attribute processing (adds rel="sponsored nofollow")
- Conversion to WordPress Block Editor (Gutenberg) format via blocks module

Migrated from: ExampleProject/tools/wordpress/docx_utils.py
Updated location: cli-tools/wordpress/wordpress_cli/utils/docx.py
"""

from __future__ import annotations

import uuid
from io import BytesIO
from pathlib import Path
from typing import Callable, Tuple

import mammoth
from bs4 import BeautifulSoup
from docx import Document

from .blocks import convert_html_to_blocks


class DocxConversionError(RuntimeError):
    pass


def extract_post_from_docx(
    word_doc_path: Path,
    upload_image: Callable[[str, bytes, str], str],
) -> Tuple[str, str]:
    if not word_doc_path.exists():
        raise FileNotFoundError(f"DOCX file not found at {word_doc_path}")

    title = _extract_title(word_doc_path)
    if not title:
        raise DocxConversionError("No heading found to use as a post title")

    html = _convert_docx_to_html(word_doc_path, upload_image)
    return title, html


def _extract_title(word_doc_path: Path) -> str:
    document = Document(word_doc_path)
    for paragraph in document.paragraphs:
        style = paragraph.style
        if style is None:
            continue
        style_name = style.name or ""
        if style_name.startswith("Heading"):
            text = paragraph.text.strip()
            if text:
                return text
    return ""


def _convert_docx_to_html(
    word_doc_path: Path,
    upload_image: Callable[[str, bytes, str], str],
) -> str:
    document = Document(word_doc_path)
    _remove_first_heading(document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    style_map = _get_style_map()

    result = mammoth.convert_to_html(
        buffer,
        style_map=style_map,
        convert_image=mammoth.images.inline(
            lambda image: _handle_image(word_doc_path, image, upload_image)
        ),
    )

    if result.messages:
        error_messages = [msg for msg in result.messages if "error" in str(msg).lower()]
        if error_messages:
            messages = ", ".join(str(msg) for msg in error_messages)
            raise DocxConversionError(f"DOCX conversion produced errors: {messages}")

    html = result.value
    html = _add_link_attributes(html)
    html = convert_html_to_blocks(html)
    return html


def _get_style_map() -> str:
    """Return style map for converting Word styles to HTML.

    Note: List Paragraph is intentionally omitted to allow mammoth to
    convert Word lists to proper HTML <ul> and <ol> elements.
    """
    return """
        p[style-name='Body Text'] => p:fresh
        p[style-name='Normal'] => p:fresh
        p[style-name='Quote'] => blockquote:fresh
        p[style-name='Intense Quote'] => blockquote:fresh
        p[style-name='Caption'] => p.caption:fresh
        p[style-name='Subtitle'] => p.subtitle:fresh
    """


def _remove_first_heading(document: Document) -> None:
    for paragraph in document.paragraphs:
        style = paragraph.style
        if style is None:
            continue
        style_name = style.name or ""
        if style_name.startswith("Heading") and paragraph.text.strip():
            element = paragraph._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
            break


def _add_link_attributes(html: str) -> str:
    """Add rel='sponsored nofollow' to all links in the HTML."""
    soup = BeautifulSoup(html, "html.parser")

    for link in soup.find_all("a"):
        existing_rel = link.get("rel", [])
        if isinstance(existing_rel, str):
            existing_rel = [existing_rel]

        new_rel = set(existing_rel)
        new_rel.add("sponsored")
        new_rel.add("nofollow")

        link["rel"] = " ".join(sorted(new_rel))

    return str(soup)


def _handle_image(
    word_doc_path: Path,
    image: mammoth.images.Image,
    upload_image: Callable[[str, bytes, str], str],
) -> dict:
    content_type = image.content_type
    extension = _extension_for_content_type(content_type)
    filename = f"{word_doc_path.stem}-{uuid.uuid4().hex[:8]}.{extension}"

    with image.open() as image_bytes:
        data = image_bytes.read()

    url = upload_image(filename, data, content_type)
    alt_text = image.alt_text or ""
    return {"src": url, "alt": alt_text}


def _extension_for_content_type(content_type: str) -> str:
    mapping = {
        "image/png": "png",
        "image/jpeg": "jpg",
        "image/jpg": "jpg",
        "image/gif": "gif",
        "image/webp": "webp",
        "image/svg+xml": "svg",
    }
    if content_type not in mapping:
        raise DocxConversionError(f"Unsupported image content type: {content_type}")
    return mapping[content_type]
