"""Regression tests for MswordClient.apply_tracked_edits (docs edit-tracked).

Covers: plain-document batch edits produce well-formed w:ins/w:del pairs;
pre-existing comments AND pre-existing tracked changes survive a new batch
edit byte-for-byte; a batch containing an unmatched old_text raises
ClientError and leaves the file on disk unmodified; and multiple occurrences
of the same text within one batch each target the correct instance.
"""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import docx
import pytest

from cli_tools_shared.exceptions import ClientError
from msword_cli.client import NSMAP, MswordClient

FIXTURE = Path(__file__).parent / "fixtures" / "tracked-changes-sample.docx"
HYPERLINK_FIXTURE = Path(__file__).parent / "fixtures" / "hyperlink-sample.docx"
W = NSMAP["w"]


def _make_doc(tmp_path: Path, name: str, text: str) -> Path:
    """Build a plain .docx with a single paragraph of text."""
    path = tmp_path / name
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _find_all(root, tag):
    return root.findall(f".//{{{W}}}{tag}")


# ---------------------------------------------------------------------------
# 1. Plain document, no existing markup.
# ---------------------------------------------------------------------------


def test_batch_edits_produce_tracked_changes(tmp_path):
    path = _make_doc(
        tmp_path,
        "plain.docx",
        "The quick brown fox jumps over the lazy dog near the old barn.",
    )

    client = MswordClient()
    result = client.apply_tracked_edits(
        str(path),
        [
            {"old_text": "quick brown fox", "new_text": "swift red fox"},
            {"old_text": "lazy dog", "new_text": "sleepy cat"},
            {"old_text": "old barn", "new_text": "new shed"},
        ],
        author="Adam Bertram",
    )

    assert result.file == str(path)
    assert result.author == "Adam Bertram"
    assert result.edits_applied == 3
    assert [e.old_text for e in result.edits] == ["quick brown fox", "lazy dog", "old barn"]
    assert [e.new_text for e in result.edits] == ["swift red fox", "sleepy cat", "new shed"]

    # Every del/ins id is unique across the whole batch.
    ids = []
    for edit in result.edits:
        ids.extend([edit.del_id, edit.ins_id])
    assert len(ids) == len(set(ids))

    root = ET.fromstring(_document_xml(path))
    del_els = _find_all(root, "del")
    ins_els = _find_all(root, "ins")
    assert len(del_els) == 3
    assert len(ins_els) == 3

    for del_el, expected_old, expected_id in zip(
        del_els, ["quick brown fox", "lazy dog", "old barn"], [e.del_id for e in result.edits]
    ):
        assert del_el.get(f"{{{W}}}id") == expected_id
        assert del_el.get(f"{{{W}}}author") == "Adam Bertram"
        assert del_el.get(f"{{{W}}}date", "").endswith("Z")
        del_text_els = del_el.findall(f".//{{{W}}}delText")
        assert "".join(t.text or "" for t in del_text_els) == expected_old
        # No w:t should remain inside a w:del.
        assert del_el.findall(f".//{{{W}}}t") == []

    for ins_el, expected_new, expected_id in zip(
        ins_els, ["swift red fox", "sleepy cat", "new shed"], [e.ins_id for e in result.edits]
    ):
        assert ins_el.get(f"{{{W}}}id") == expected_id
        assert ins_el.get(f"{{{W}}}author") == "Adam Bertram"
        t_els = ins_el.findall(f".//{{{W}}}t")
        assert "".join(t.text or "" for t in t_els) == expected_new

    # The document still opens cleanly with python-docx afterward.
    reopened = docx.Document(str(path))
    assert reopened.paragraphs


# ---------------------------------------------------------------------------
# 2. Pre-existing comment AND pre-existing tracked change must survive.
# ---------------------------------------------------------------------------


def test_preexisting_comment_and_tracked_change_survive_new_edits(tmp_path):
    path = tmp_path / "existing-markup.docx"
    shutil.copyfile(FIXTURE, path)

    client = MswordClient()
    before_comments = client.extract_comments(str(path))
    assert len(before_comments) == 1
    assert before_comments[0].id == "5"

    before_root = ET.fromstring(_document_xml(path))
    before_del = _find_all(before_root, "del")[0]
    before_ins = _find_all(before_root, "ins")[0]
    assert before_del.get(f"{{{W}}}id") == "10"
    assert before_ins.get(f"{{{W}}}id") == "11"

    result = client.apply_tracked_edits(
        str(path),
        [{"old_text": "updated roadmap", "new_text": "revised roadmap"}],
        author="Adam Bertram",
    )
    assert result.edits_applied == 1
    # New ids must not collide with the existing comment (5) or tracked-change (10, 11) ids.
    assert result.edits[0].del_id not in {"5", "10", "11"}
    assert result.edits[0].ins_id not in {"5", "10", "11"}

    after_comments = client.extract_comments(str(path))
    assert after_comments == before_comments

    after_root = ET.fromstring(_document_xml(path))
    after_del = next(e for e in _find_all(after_root, "del") if e.get(f"{{{W}}}id") == "10")
    after_ins = next(e for e in _find_all(after_root, "ins") if e.get(f"{{{W}}}id") == "11")

    assert after_del.get(f"{{{W}}}author") == "Prior Author"
    assert after_ins.get(f"{{{W}}}author") == "Prior Author"
    assert "".join(t.text or "" for t in after_del.findall(f".//{{{W}}}delText")) == "old goal"
    assert "".join(t.text or "" for t in after_ins.findall(f".//{{{W}}}t")) == "new goal"

    # New del/ins pair exists for the new edit, attributed to the new author.
    new_del = next(e for e in _find_all(after_root, "del") if e.get(f"{{{W}}}id") == result.edits[0].del_id)
    new_ins = next(e for e in _find_all(after_root, "ins") if e.get(f"{{{W}}}id") == result.edits[0].ins_id)
    assert new_del.get(f"{{{W}}}author") == "Adam Bertram"
    assert new_ins.get(f"{{{W}}}author") == "Adam Bertram"

    reopened = docx.Document(str(path))
    assert reopened.paragraphs


# ---------------------------------------------------------------------------
# 3. Unmatched old_text raises ClientError and does not modify the file.
# ---------------------------------------------------------------------------


def test_unmatched_old_text_raises_and_leaves_file_unmodified(tmp_path):
    path = _make_doc(tmp_path, "unmatched.docx", "Alpha beta gamma delta.")
    original_bytes = path.read_bytes()

    client = MswordClient()
    with pytest.raises(ClientError, match="Reference text not found"):
        client.apply_tracked_edits(
            str(path),
            [
                {"old_text": "alpha beta", "new_text": "x"},
                {"old_text": "does-not-exist-anywhere", "new_text": "y"},
            ],
            author="Adam Bertram",
        )

    assert path.read_bytes() == original_bytes


# ---------------------------------------------------------------------------
# 4. Multiple occurrences of the same text, different occurrence values.
# ---------------------------------------------------------------------------


def test_multiple_occurrences_target_correct_instance(tmp_path):
    path = _make_doc(
        tmp_path,
        "occurrences.docx",
        "widget one. widget two. widget three.",
    )

    client = MswordClient()
    # Editing highest occurrence first keeps earlier occurrence numbers stable
    # for the edits that follow in the same batch (see apply_tracked_edits docstring).
    result = client.apply_tracked_edits(
        str(path),
        [
            {"old_text": "widget", "new_text": "gadget", "occurrence": 3},
            {"old_text": "widget", "new_text": "gizmo", "occurrence": 1},
        ],
        author="Adam Bertram",
    )
    assert result.edits_applied == 2

    root = ET.fromstring(_document_xml(path))
    ins_els = _find_all(root, "ins")
    # Elements appear in document order (occurrence 1 precedes occurrence 3),
    # regardless of the order edits were listed/applied in the batch.
    ins_texts = ["".join(t.text or "" for t in el.findall(f".//{{{W}}}t")) for el in ins_els]
    assert ins_texts == ["gizmo", "gadget"]

    del_els = _find_all(root, "del")
    assert len(del_els) == 2
    for el in del_els:
        assert "".join(t.text or "" for t in el.findall(f".//{{{W}}}delText")) == "widget"

    # "widget two" (the untouched middle occurrence) is still plain, unedited text.
    full_text = "".join(t.text or "" for t in _find_all(root, "t"))
    assert "widget two" in full_text


# ---------------------------------------------------------------------------
# 5. old_text spanning a w:hyperlink must not be mistaken for a comment or
#    tracked-change overlap (regression for the false-positive refusal).
# ---------------------------------------------------------------------------


def test_old_text_spanning_hyperlink_is_not_a_false_positive(tmp_path):
    path = tmp_path / "hyperlink.docx"
    shutil.copyfile(HYPERLINK_FIXTURE, path)

    client = MswordClient()
    before_comments = client.extract_comments(str(path))
    assert {c.id for c in before_comments} == {"99", "77"}

    # old_text starts in a plain run, crosses entirely through the
    # w:hyperlink's own nested w:r, and ends in a trailing plain run -- no
    # comment or tracked change is anywhere near this paragraph.
    result = client.apply_tracked_edits(
        str(path),
        [
            {
                "old_text": "Before text link anchor text after text.",
                "new_text": "Replacement text.",
            }
        ],
        author="Adam Bertram",
    )
    assert result.edits_applied == 1

    # The unrelated comment (99) and the genuinely-overlapping one (77, on a
    # different paragraph this edit never touched) both survive untouched.
    after_comments = client.extract_comments(str(path))
    assert after_comments == before_comments

    root = ET.fromstring(_document_xml(path))
    del_el = _find_all(root, "del")[0]
    ins_el = _find_all(root, "ins")[0]
    del_text = "".join(t.text or "" for t in del_el.findall(f".//{{{W}}}delText"))
    assert del_text == "Before text link anchor text after text."
    ins_text = "".join(t.text or "" for t in ins_el.findall(f".//{{{W}}}t"))
    assert ins_text == "Replacement text."

    # The now-fully-consumed w:hyperlink wrapper must not remain as a dangling
    # empty element inside the w:del.
    assert del_el.findall(f".//{{{W}}}hyperlink") == []

    reopened = docx.Document(str(path))
    assert reopened.paragraphs


# ---------------------------------------------------------------------------
# 6. A genuine comment overlap must still be refused (the hyperlink fix must
#    not weaken the real protection this error exists for).
# ---------------------------------------------------------------------------


def test_true_comment_overlap_still_raises(tmp_path):
    path = tmp_path / "hyperlink.docx"
    shutil.copyfile(HYPERLINK_FIXTURE, path)
    original_bytes = path.read_bytes()

    client = MswordClient()
    with pytest.raises(
        ClientError,
        match="an existing comment or tracked-change marker sits inside the matched text",
    ):
        client.apply_tracked_edits(
            str(path),
            [{"old_text": "alpha beta gamma.", "new_text": "x"}],
            author="Adam Bertram",
        )

    assert path.read_bytes() == original_bytes


# ---------------------------------------------------------------------------
# 7. new_text containing paragraph breaks ("\n\n") must produce real w:p
#    elements per paragraph, not literal newlines inside one w:t (regression
#    for the multi-paragraph collapse bug).
# ---------------------------------------------------------------------------


def _pPr_mark_ins_id(p_el):
    """Return the w:id of p_el's own inserted-paragraph-mark w:ins, or None."""
    ins_el = p_el.find(f"{{{W}}}pPr/{{{W}}}rPr/{{{W}}}ins")
    return ins_el.get(f"{{{W}}}id") if ins_el is not None else None


def test_multi_paragraph_new_text_splits_into_separate_paragraphs(tmp_path):
    path = _make_doc(
        tmp_path,
        "multipara.docx",
        "Old placeholder text that will be replaced entirely.",
    )

    client = MswordClient()
    result = client.apply_tracked_edits(
        str(path),
        [
            {
                "old_text": "Old placeholder text that will be replaced entirely.",
                "new_text": "Summary\n\nBody paragraph one.\n\nBody paragraph two.",
            }
        ],
        author="Adam Bertram",
    )
    assert result.edits_applied == 1
    assert result.edits[0].del_id
    assert result.edits[0].ins_id
    # 2 extra content w:ins (paragraphs 2 and 3) + 2 paragraph-mark w:ins
    # (closing paragraphs 1 and 2) = 4 extra ids beyond the first ins_id.
    assert len(result.edits[0].extra_ins_ids) == 4
    all_ids = [result.edits[0].del_id, result.edits[0].ins_id] + result.edits[0].extra_ins_ids
    assert len(all_ids) == len(set(all_ids))

    root = ET.fromstring(_document_xml(path))
    p_els = _find_all(root, "p")
    assert len(p_els) == 3

    # Every w:t (not w:delText) across the three paragraphs, in document order.
    ins_texts = [
        "".join(t.text or "" for t in p.findall(f".//{{{W}}}ins//{{{W}}}t"))
        for p in p_els
    ]
    assert ins_texts == ["Summary", "Body paragraph one.", "Body paragraph two."]

    # The del element (old text) lives in the first paragraph only.
    assert _find_all(p_els[0], "del") != []
    assert _find_all(p_els[1], "del") == []
    assert _find_all(p_els[2], "del") == []

    # Paragraphs 1 and 2 (all but the last) carry a tracked paragraph-mark
    # insertion; the last paragraph keeps its own (unflagged) mark.
    assert _pPr_mark_ins_id(p_els[0]) is not None
    assert _pPr_mark_ins_id(p_els[1]) is not None
    assert _pPr_mark_ins_id(p_els[2]) is None

    # No literal "\n\n" survives anywhere in the saved document.
    assert "\n\n" not in _document_xml(path)

    reopened = docx.Document(str(path))
    assert len(reopened.paragraphs) == 3


def test_multi_paragraph_new_text_preserves_trailing_original_text(tmp_path):
    path = _make_doc(
        tmp_path,
        "multipara-trailing.docx",
        "Keep before. MATCH ME suffix stays after.",
    )

    client = MswordClient()
    client.apply_tracked_edits(
        str(path),
        [{"old_text": "MATCH ME", "new_text": "Chunk one.\n\nChunk two."}],
        author="Adam Bertram",
    )

    root = ET.fromstring(_document_xml(path))
    p_els = _find_all(root, "p")
    assert len(p_els) == 2

    # Untouched text before the match stays plain, in the first paragraph.
    first_plain = "".join(
        t.text or "" for t in p_els[0].findall(f"{{{W}}}r/{{{W}}}t")
    )
    assert first_plain == "Keep before. "

    # Untouched trailing text moves into the final paragraph (the one that
    # keeps the original, unflagged paragraph mark), after the last inserted
    # chunk -- not lost, and not left behind in the first paragraph.
    assert _pPr_mark_ins_id(p_els[1]) is None
    last_plain = "".join(
        t.text or "" for t in p_els[1].findall(f"{{{W}}}r/{{{W}}}t")
    )
    assert last_plain == " suffix stays after."

    reopened = docx.Document(str(path))
    assert len(reopened.paragraphs) == 2
