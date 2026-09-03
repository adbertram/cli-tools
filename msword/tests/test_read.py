"""Regression tests for MswordClient.read_document (docs read).

Covers: plain-paragraph reads still work; a Word table's cell text is
included (doc.paragraphs, python-docx's high-level accessor, silently skips
every paragraph nested inside a w:tbl cell, so read_document previously
dropped table content with no warning); and a table's rows appear in their
original document position relative to the surrounding paragraphs.
"""
from __future__ import annotations

from pathlib import Path

import docx

from msword_cli.client import MswordClient


def _make_doc(tmp_path: Path, name: str, paragraphs: list) -> Path:
    """Build a plain .docx with one paragraph per entry in `paragraphs`."""
    path = tmp_path / name
    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)
    return path


def test_read_document_plain_paragraphs(tmp_path):
    path = _make_doc(
        tmp_path,
        "plain.docx",
        ["First paragraph.", "Second paragraph.", "Third paragraph."],
    )

    client = MswordClient()
    result = client.read_document(str(path))

    assert result.paragraphs == 3
    assert result.content == "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."


def test_read_document_includes_table_content(tmp_path):
    path = tmp_path / "with-table.docx"
    doc = docx.Document()
    doc.add_paragraph("Intro paragraph before the table.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Lifecycle stage"
    table.cell(0, 1).text = "Automated before AI"
    table.cell(1, 0).text = "Test maintenance"
    table.cell(1, 1).text = "Retry-on-red, quarantine lists"
    doc.add_paragraph("Outro paragraph after the table.")
    doc.save(path)

    client = MswordClient()
    result = client.read_document(str(path))

    # Every cell's text made it into the output -- none of this reached
    # read_document before the fix, since doc.paragraphs skips w:tbl cells.
    for expected in (
        "Lifecycle stage",
        "Automated before AI",
        "Test maintenance",
        "Retry-on-red, quarantine lists",
    ):
        assert expected in result.content

    # The table's rows stay between the surrounding paragraphs, in the
    # table's original document position, not appended at the end.
    entries = result.content.split("\n\n")
    assert entries[0] == "Intro paragraph before the table."
    assert entries[1] == "Lifecycle stage | Automated before AI"
    assert entries[2] == "Test maintenance | Retry-on-red, quarantine lists"
    assert entries[3] == "Outro paragraph after the table."
    assert result.paragraphs == 4
