"""Msword client for reading, converting, and extracting comments from Word documents."""
import copy
import os
from datetime import datetime, timezone
from typing import Dict, List
from xml.etree import ElementTree as ET

import docx
import mammoth
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree as LET

from .models import (
    Comment,
    DocumentContent,
    ConvertedDocument,
    AddCommentResult,
    TrackedEditApplied,
    EditTrackedChangesResult,
)

from cli_tools_shared.exceptions import ClientError

# Word XML namespaces
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
}

for _prefix, _uri in NSMAP.items():
    ET.register_namespace(_prefix, _uri)
ET.register_namespace('r', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships')

COMMENTS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"


def _find_comments_part(doc):
    """Return the word/comments.xml part, or None when the document has no comments.

    Matches the relationship type exactly. A substring test such as
    ``"comments" in rel.reltype`` also matches the sibling parts
    ``commentsIds``, ``commentsExtended``, and ``commentsExtensible``; those
    parts hold metadata, not ``w:comment`` elements, so selecting one of them
    yields zero comments.
    """
    for rel in doc.part.rels.values():
        if getattr(rel, "is_external", False):
            continue
        if rel.reltype == COMMENTS_REL_TYPE:
            return rel.target_part
    return None


def _validate_file(file_path: str) -> str:
    """Validate that the file exists and is a .docx file."""
    path = os.path.expanduser(file_path)
    if not os.path.isfile(path):
        raise ClientError(f"File not found: {path}")
    if not path.lower().endswith(".docx"):
        raise ClientError(f"Not a Word document (.docx): {path}")
    return path


def _local_tag(el) -> str:
    """Return an element's tag without its namespace prefix."""
    return el.tag.split("}")[-1] if "}" in el.tag else el.tag


class MswordClient:
    """Client for processing Word documents."""

    def read_document(self, file_path: str) -> DocumentContent:
        """Read text content from a Word document.

        Args:
            file_path: Path to the .docx file

        Returns:
            DocumentContent model with full text
        """
        path = _validate_file(file_path)
        doc = docx.Document(path)

        paragraphs = []
        for text in self._iter_body_text(doc):
            text = text.strip()
            if text:
                paragraphs.append(text)

        return DocumentContent(
            file=path,
            paragraphs=len(paragraphs),
            content="\n\n".join(paragraphs),
        )

    def _iter_body_text(self, doc):
        """Yield paragraph text and table-row text, in document order.

        ``doc.paragraphs`` (python-docx's high-level accessor) returns only
        the body's direct ``w:p`` children and silently skips every
        paragraph nested inside a ``w:tbl`` cell, so a document's table
        content never reached ``read_document`` callers. Walk the body's
        direct children instead: a paragraph yields its own text; a table
        yields one entry per row, with that row's cells joined by " | ",
        in the table's original position among the surrounding paragraphs.
        Any other direct-child element (e.g. ``w:sectPr``, a bookmark) is
        skipped, matching ``doc.paragraphs``' existing behavior for those.
        """
        for child in doc.element.body:
            tag = _local_tag(child)
            if tag == "p":
                yield Paragraph(child, doc).text
            elif tag == "tbl":
                table = Table(child, doc)
                for row in table.rows:
                    yield " | ".join(cell.text.strip() for cell in row.cells)

    def convert_to_markdown(self, file_path: str) -> ConvertedDocument:
        """Convert a Word document to Markdown.

        Args:
            file_path: Path to the .docx file

        Returns:
            ConvertedDocument model with markdown content
        """
        path = _validate_file(file_path)

        with open(path, "rb") as f:
            result = mammoth.convert_to_markdown(f)

        return ConvertedDocument(
            file=path,
            markdown=result.value,
            messages=[str(m) for m in result.messages],
        )

    def extract_comments(self, file_path: str) -> List[Comment]:
        """Extract comments with their referenced context from a Word document.

        Parses the document XML to find comment markers and extracts
        the text between commentRangeStart and commentRangeEnd elements.

        Args:
            file_path: Path to the .docx file

        Returns:
            List of Comment models with context
        """
        path = _validate_file(file_path)
        doc = docx.Document(path)

        # Extract comments from the word/comments.xml part (exact reltype match).
        comments_part = _find_comments_part(doc)
        if comments_part is None:
            return []

        # Parse comments XML. Both classic comments (small integer w:id) and
        # modern threaded comments (large integer w:id, w:date without a
        # trailing "Z", an extra xmlns:w redeclaration on the element) live in
        # this part and are read identically.
        comments_xml = ET.fromstring(comments_part.blob)
        comment_data = {}
        for comment_el in comments_xml.findall(qn("w:comment")):
            cid = comment_el.get(qn("w:id"))
            author = comment_el.get(qn("w:author"), "Unknown")
            date = comment_el.get(qn("w:date"))

            # Get comment text from all paragraphs/runs.
            texts = []
            for t_el in comment_el.iter(qn("w:t")):
                if t_el.text:
                    texts.append(t_el.text)

            comment_data[cid] = {
                "id": cid,
                "author": author,
                "date": date,
                "text": " ".join(texts),
            }

        # Extract context for each comment from the document body
        body_xml = doc.element.body
        context_map = self._extract_comment_contexts(body_xml)

        # Build Comment models
        comments = []
        for cid, data in comment_data.items():
            context = context_map.get(cid)
            comments.append(
                Comment(
                    id=data["id"],
                    author=data["author"],
                    date=data["date"],
                    text=data["text"],
                    context=context,
                )
            )

        return comments

    def get_comment(self, file_path: str, comment_id: str) -> Comment:
        """Return a single comment by ID from a Word document."""
        comments = self.extract_comments(file_path)
        for comment in comments:
            if comment.id == comment_id:
                return comment
        raise ClientError(f"Comment not found: {comment_id}")

    def add_comment(
        self, file_path: str, text: str, author: str, reference_text: str, occurrence: int = 1
    ) -> AddCommentResult:
        """Add an inline comment anchored to specific text in a Word document."""
        if not reference_text:
            raise ClientError("reference_text cannot be empty")
        if not text.strip():
            raise ClientError("Comment text cannot be empty")
        if not author.strip():
            raise ClientError("Author cannot be empty")
        if occurrence < 1:
            raise ClientError("occurrence must be a positive integer")

        path = _validate_file(file_path)
        doc = docx.Document(path)

        start_el, end_el = self._find_reference_text(doc, reference_text, occurrence)

        comments_part = self._get_comments_part(doc)
        next_id = self._get_next_comment_id(doc, comments_part)

        self._add_comment_xml(comments_part, next_id, text, author)
        self._insert_comment_markers(start_el, end_el, next_id)

        doc.save(path)

        return AddCommentResult(
            file=path,
            comment_id=str(next_id),
            author=author,
            text=text,
            reference_text=reference_text,
        )

    def apply_tracked_edits(
        self, file_path: str, edits: List[dict], author: str
    ) -> EditTrackedChangesResult:
        """Apply a batch of text edits as Word tracked changes (w:ins/w:del).

        Each edit locates ``old_text`` (via ``_find_reference_text``, reusing
        the same run-splitting logic as ``add_comment``), wraps the isolated
        run(s) in a ``w:del`` (converting ``w:t`` to ``w:delText``), and
        inserts a sibling ``w:ins`` containing ``new_text``, attributed to
        ``author``. Every pre-existing comment and tracked change is verified
        byte-for-byte unchanged before the file is saved; if anything from the
        original document is missing or altered, the file is left untouched
        and a ``ClientError`` is raised.

        Edits are applied in list order against the document as it stands
        after prior edits in the same batch. ``occurrence`` is resolved
        against that current state, so if two edits share the same
        ``old_text``, order edits from the highest occurrence to the lowest
        to avoid occurrence numbers shifting after an earlier match in the
        same text is replaced.
        """
        if not author.strip():
            raise ClientError("Author cannot be empty")
        if not edits:
            raise ClientError("edits cannot be empty")
        for index, edit in enumerate(edits):
            if not edit.get("old_text", "").strip():
                raise ClientError(f"edits[{index}].old_text cannot be empty")
            if not isinstance(edit.get("new_text"), str):
                raise ClientError(f"edits[{index}].new_text is required")
            if edit.get("occurrence", 1) < 1:
                raise ClientError(f"edits[{index}].occurrence must be a positive integer")

        path = _validate_file(file_path)
        doc = docx.Document(path)

        baseline = self._capture_markup_baseline(doc)

        comments_part = _find_comments_part(doc)
        next_id = self._get_next_change_id(doc, comments_part)

        date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        applied: List[TrackedEditApplied] = []

        for edit in edits:
            old_text = edit["old_text"]
            new_text = edit["new_text"]
            occurrence = edit.get("occurrence", 1)

            start_el, end_el = self._find_reference_text(doc, old_text, occurrence)
            run_range = self._collect_run_range(start_el, end_el)

            del_id = next_id
            next_id += 1

            del_el = self._build_change_element("del", del_id, author, date)
            start_el.addprevious(del_el)
            for run_el in run_range:
                self._convert_run_to_del_text(run_el)
                del_el.append(run_el)

            source_run = run_range[0]
            paragraph_chunks = new_text.split("\n\n")
            cur_p = del_el.getparent()
            anchor = del_el
            ins_ids = []

            for chunk_index, chunk in enumerate(paragraph_chunks):
                ins_id = next_id
                next_id += 1
                ins_ids.append(ins_id)

                ins_el = self._build_change_element("ins", ins_id, author, date)
                ins_el.append(self._build_ins_run(source_run, chunk))
                if anchor is None:
                    cur_p.insert(0, ins_el)
                else:
                    anchor.addnext(ins_el)
                anchor = ins_el

                if chunk_index < len(paragraph_chunks) - 1:
                    # new_text carries another paragraph break here: split the
                    # host paragraph so the next chunk becomes its own w:p,
                    # with an inserted (tracked) paragraph mark closing this one.
                    mark_id = next_id
                    next_id += 1
                    ins_ids.append(mark_id)
                    cur_p, anchor = self._split_paragraph_for_inserted_break(
                        cur_p, anchor, author, date, mark_id
                    )

            applied.append(
                TrackedEditApplied(
                    old_text=old_text,
                    new_text=new_text,
                    occurrence=occurrence,
                    del_id=str(del_id),
                    ins_id=str(ins_ids[0]),
                    extra_ins_ids=[str(i) for i in ins_ids[1:]],
                )
            )

        self._verify_markup_baseline(doc, baseline)

        doc.save(path)

        return EditTrackedChangesResult(
            file=path,
            author=author,
            edits_applied=len(applied),
            edits=applied,
        )

    # Elements that wrap one or more w:r runs as their own children rather
    # than exposing them as direct paragraph-level siblings. A run inside one
    # of these is still plain matchable/deletable text; it just needs to be
    # promoted to a paragraph-level sibling before the sibling-only walk in
    # _collect_run_range can reach it.
    _RUN_WRAPPER_TAGS = ("hyperlink",)

    def _collect_run_range(self, start_el, end_el) -> list:
        """Collect the sibling w:r run elements from start_el to end_el inclusive.

        Fails loudly rather than silently skipping over a non-run sibling
        (such as an existing comment marker) that falls inside the range,
        since blindly wrapping it in w:del/w:ins would corrupt it. Runs
        nested inside a wrapper element (e.g. w:hyperlink) are promoted to
        paragraph-level siblings first (see _promote_wrapped_runs) so a match
        that starts or ends outside a hyperlink but crosses through it is not
        mistaken for overlapping a real comment or tracked-change marker.
        """
        self._promote_wrapped_runs(start_el, end_el)

        runs = []
        current = start_el
        while True:
            tag = _local_tag(current)
            if tag != "r":
                raise ClientError(
                    "Cannot apply tracked-change edit: an existing comment or "
                    "tracked-change marker sits inside the matched text"
                )
            runs.append(current)
            if current is end_el:
                return runs
            current = current.getnext()
            if current is None:
                raise ClientError("Failed to collect run range for tracked-change edit")

    def _promote_wrapped_runs(self, start_el, end_el):
        """Promote w:r runs inside a wrapper element (e.g. w:hyperlink) within
        [start_el, end_el] to direct paragraph-level siblings, in place of the
        wrapper, so the sibling-only walk in _collect_run_range can reach them.

        The overall match range is contiguous in document order, and a
        wrapper's own children are contiguous too, so the in-range subset of
        any one wrapper's children is always a contiguous prefix, a suffix,
        the whole set, or (when the match starts and ends inside the same
        wrapper without touching either of its own edges) a self-contained
        interior slice that the plain sibling walk already handles without
        promotion — that last case is left untouched here.
        """
        para_el = start_el.getparent()
        while para_el is not None and _local_tag(para_el) != "p":
            para_el = para_el.getparent()
        if para_el is None:
            raise ClientError("Failed to locate enclosing paragraph for tracked-change edit")

        run_els = list(para_el.iter(qn("w:r")))
        try:
            start_idx = run_els.index(start_el)
            end_idx = run_els.index(end_el)
        except ValueError:
            raise ClientError("Failed to collect run range for tracked-change edit")

        matched_ids = {id(r) for r in run_els[start_idx:end_idx + 1]}

        wrappers = []
        for run_el in run_els[start_idx:end_idx + 1]:
            parent = run_el.getparent()
            if (
                parent is not None
                and _local_tag(parent) in self._RUN_WRAPPER_TAGS
                and parent not in wrappers
            ):
                wrappers.append(parent)

        for wrapper in wrappers:
            wrapper_runs = [r for r in wrapper if _local_tag(r) == "r"]
            in_range = [i for i, r in enumerate(wrapper_runs) if id(r) in matched_ids]
            lo, hi = in_range[0], in_range[-1]
            if lo != 0 and hi != len(wrapper_runs) - 1:
                # Self-contained interior match: no boundary crossed, no
                # promotion needed for this wrapper.
                continue
            if lo == 0:
                for i in range(lo, hi + 1):
                    wrapper.addprevious(wrapper_runs[i])
            else:
                for i in range(hi, lo - 1, -1):
                    wrapper.addnext(wrapper_runs[i])
            if wrapper.find(qn("w:r")) is None:
                wrapper.getparent().remove(wrapper)

    def _build_change_element(self, tag: str, change_id: int, author: str, date: str):
        """Build an empty w:ins or w:del element with id/author/date attributes."""
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:id"), str(change_id))
        el.set(qn("w:author"), author)
        el.set(qn("w:date"), date)
        return el

    def _convert_run_to_del_text(self, run_el):
        """Convert a run's w:t children to w:delText in place (OOXML tracked-deletion convention)."""
        for t_el in list(run_el.findall(qn("w:t"))):
            del_text_el = OxmlElement("w:delText")
            del_text_el.text = t_el.text
            space = t_el.get("{http://www.w3.org/XML/1998/namespace}space")
            if space:
                del_text_el.set("{http://www.w3.org/XML/1998/namespace}space", space)
            run_el.replace(t_el, del_text_el)

    def _build_ins_run(self, source_run_el, new_text: str):
        """Build a new w:r/w:t run for a w:ins, cloning rPr from source_run_el when present."""
        new_run = OxmlElement("w:r")
        rPr = source_run_el.find(qn("w:rPr"))
        if rPr is not None:
            new_run.append(copy.deepcopy(rPr))
        t_el = OxmlElement("w:t")
        t_el.text = new_text
        if new_text and (new_text[0] == " " or new_text[-1] == " "):
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_run.append(t_el)
        return new_run

    def _split_paragraph_for_inserted_break(self, cur_p, after_el, author: str, date: str, mark_id: int):
        """Split cur_p into two adjacent w:p elements right after after_el.

        A "\\n\\n" inside new_text means the replacement content spans more
        than one Word paragraph, which OOXML represents as separate w:p
        elements -- not literal newline characters inside a single w:t
        (Word never treats an embedded newline in w:t as a paragraph break).

        Everything that follows after_el in cur_p (any untouched trailing
        content from the original paragraph) moves into a new sibling w:p,
        cloning cur_p's own w:pPr so the trailing paragraph keeps the same
        style/list formatting. cur_p's own paragraph mark is then flagged as
        a tracked insertion (w:pPr/w:rPr/w:ins) since this break did not
        exist in the original document -- the OOXML convention for a newly
        inserted paragraph mark. Any revision marker the ORIGINAL paragraph
        mark already carried (ins/del/moveFrom/moveTo) describes that
        original mark, which now belongs to the new trailing paragraph, not
        to this freshly created break, so it is left on the clone and
        stripped from cur_p's copy.

        Returns (new_p, insertion_anchor), where insertion_anchor is the
        element the next chunk's w:ins should be inserted after via
        addnext(), or None when new_p has no w:pPr to anchor on (the next
        chunk must then be inserted with new_p.insert(0, ...)).
        """
        new_p = OxmlElement("w:p")
        cur_pPr = cur_p.find(qn("w:pPr"))
        new_pPr = None
        if cur_pPr is not None:
            new_pPr = copy.deepcopy(cur_pPr)
            new_p.append(new_pPr)

        moving = []
        sib = after_el.getnext()
        while sib is not None:
            moving.append(sib)
            sib = sib.getnext()
        for el in moving:
            new_p.append(el)

        cur_p.addnext(new_p)

        if cur_pPr is None:
            cur_pPr = OxmlElement("w:pPr")
            cur_p.insert(0, cur_pPr)
        cur_rPr = cur_pPr.find(qn("w:rPr"))
        if cur_rPr is None:
            cur_rPr = OxmlElement("w:rPr")
            cur_pPr.append(cur_rPr)
        else:
            for prior_tag in ("ins", "del", "moveFrom", "moveTo"):
                prior = cur_rPr.find(qn(f"w:{prior_tag}"))
                if prior is not None:
                    cur_rPr.remove(prior)
        mark_ins = OxmlElement("w:ins")
        mark_ins.set(qn("w:id"), str(mark_id))
        mark_ins.set(qn("w:author"), author)
        mark_ins.set(qn("w:date"), date)
        cur_rPr.insert(0, mark_ins)

        return new_p, new_pPr

    def _comment_ids(self, comments_part) -> set:
        """Return the set of w:id values from every w:comment in comments.xml."""
        ids = set()
        if comments_part is not None:
            for comment_el in comments_part.element.findall(qn("w:comment")):
                cid = comment_el.get(qn("w:id"))
                if cid is not None:
                    ids.add(cid)
        return ids

    def _max_markup_id(self, doc, comments_part, extra_tags: tuple = ()) -> int:
        """Return the highest existing w:id across comments.xml and matching
        body markers, so a newly allocated id never collides with one already
        in use. extra_tags adds body element tags beyond the comment markers
        (e.g. "ins"/"del" for tracked-change id allocation)."""
        max_id = max((int(cid) for cid in self._comment_ids(comments_part)), default=0)

        tags = ("commentRangeStart", "commentRangeEnd", "commentReference") + extra_tags
        for el in doc.element.body.iter():
            if _local_tag(el) in tags:
                cid = el.get(qn("w:id"))
                if cid is not None:
                    max_id = max(max_id, int(cid))

        return max_id

    def _get_next_change_id(self, doc, comments_part) -> int:
        """Get the next available w:id, scanning comments.xml, body comment
        markers, and existing w:ins/w:del elements so new tracked-change ids
        never collide with existing comment or tracked-change ids."""
        return self._max_markup_id(doc, comments_part, extra_tags=("ins", "del")) + 1

    def _capture_markup_baseline(self, doc) -> Dict:
        """Snapshot every pre-existing comment id and tracked-change element.

        Captured immediately after opening the document, before any edit is
        applied, so it reflects the original file on disk exactly.
        """
        comments_part = _find_comments_part(doc)
        return {"comment_ids": self._comment_ids(comments_part), "changes": self._snapshot_changes(doc)}

    def _snapshot_changes(self, doc) -> Dict:
        """Map (tag, w:id) -> exact serialized XML for every existing w:ins/w:del element."""
        changes = {}
        for el in doc.element.body.iter():
            tag = _local_tag(el)
            if tag in ("ins", "del"):
                cid = el.get(qn("w:id"))
                if cid is not None:
                    changes[(tag, cid)] = LET.tostring(el, encoding="unicode")
        return changes

    def _verify_markup_baseline(self, doc, baseline: Dict):
        """Confirm every pre-existing comment and tracked change is still
        present and byte-for-byte unchanged after applying edits.

        Runs after all edits are applied but before the document is saved.
        Raises ClientError without saving if anything from the original
        document is missing or altered.
        """
        comments_part = _find_comments_part(doc)
        missing_comments = baseline["comment_ids"] - self._comment_ids(comments_part)
        if missing_comments:
            raise ClientError(
                "Verification failed: existing comment(s) missing after edit: "
                f"{sorted(missing_comments)}"
            )

        current_changes = self._snapshot_changes(doc)
        for key, original_xml in baseline["changes"].items():
            current_xml = current_changes.get(key)
            if current_xml is None:
                raise ClientError(
                    f"Verification failed: existing tracked change {key[0]}#{key[1]} "
                    "is missing after edit"
                )
            if current_xml != original_xml:
                raise ClientError(
                    f"Verification failed: existing tracked change {key[0]}#{key[1]} "
                    "was altered by this edit"
                )

    def _find_reference_text(self, doc, reference_text: str, occurrence: int):
        """Find and isolate the run elements containing the target text.

        Searches body paragraphs, table cells, headers, and footers.
        Returns (start_run_el, end_run_el) after splitting boundary runs.
        """
        count = 0
        for para_el in self._iter_paragraph_elements(doc):
            run_els = list(para_el.iter(qn("w:r")))
            if not run_els:
                continue
            runs_text = "".join(self._run_el_text(r) for r in run_els)
            start = 0
            while True:
                idx = runs_text.find(reference_text, start)
                if idx == -1:
                    break
                count += 1
                if count == occurrence:
                    end_pos = idx + len(reference_text)
                    return self._isolate_match_runs(run_els, idx, end_pos)
                start = idx + 1

        raise ClientError(
            f"Reference text not found: '{reference_text}'"
            + (f" (occurrence {occurrence})" if occurrence > 1 else "")
        )

    def _iter_paragraph_elements(self, doc):
        """Yield all w:p elements: body (including tables), headers, footers."""
        for p_el in doc.element.body.iter(qn("w:p")):
            yield p_el
        seen = set()
        for section in doc.sections:
            for part in (section.header, section.footer):
                if part is None:
                    continue
                el_id = id(part._element)
                if el_id in seen:
                    continue
                seen.add(el_id)
                for p_el in part._element.iter(qn("w:p")):
                    yield p_el

    def _run_el_text(self, run_el):
        """Get text from a w:r element's w:t children."""
        return "".join(t.text for t in run_el.findall(qn("w:t")) if t.text)

    def _isolate_match_runs(self, run_els, match_start: int, match_end: int):
        """Find which runs contain the match, split at boundaries, return (start_el, end_el).

        Handles both single-run and multi-run matches by splitting boundary runs
        so that only the matched text is inside the returned range.
        """
        current_pos = 0
        first_idx = last_idx = None
        first_local_start = last_local_end = 0

        for i, run_el in enumerate(run_els):
            text_len = len(self._run_el_text(run_el))
            if first_idx is None and current_pos + text_len > match_start:
                first_idx = i
                first_local_start = match_start - current_pos
            if current_pos + text_len >= match_end:
                last_idx = i
                last_local_end = match_end - current_pos
                break
            current_pos += text_len

        if first_idx is None or last_idx is None:
            raise ClientError("Failed to map reference text to document runs")

        first_el = run_els[first_idx]
        last_el = run_els[last_idx]

        if first_local_start > 0:
            first_text = self._run_el_text(first_el)
            before_run = self._clone_run(first_el, first_text[:first_local_start])
            first_el.addprevious(before_run)
            self._set_run_text(first_el, first_text[first_local_start:])
            if first_idx == last_idx:
                last_local_end -= first_local_start

        last_text = self._run_el_text(last_el)
        if last_local_end < len(last_text):
            after_run = self._clone_run(last_el, last_text[last_local_end:])
            last_el.addnext(after_run)
            self._set_run_text(last_el, last_text[:last_local_end])

        return first_el, last_el

    def _set_run_text(self, run_el, text: str):
        """Set run text by replacing w:t elements only, preserving w:br/w:tab/etc."""
        for t in list(run_el.findall(qn("w:t"))):
            run_el.remove(t)
        new_t = OxmlElement("w:t")
        new_t.text = text
        if text and (text[0] == " " or text[-1] == " "):
            new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        rPr = run_el.find(qn("w:rPr"))
        if rPr is not None:
            rPr.addnext(new_t)
        else:
            run_el.insert(0, new_t)

    def _clone_run(self, source_run_el, new_text: str):
        """Clone a run element with new text, preserving formatting."""
        import copy
        new_run = copy.deepcopy(source_run_el)
        self._set_run_text(new_run, new_text)
        return new_run

    def _get_comments_part(self, doc):
        """Get existing comments part or create a new one."""
        existing = _find_comments_part(doc)
        if existing is not None:
            return existing

        from docx.oxml import parse_xml

        w_ns = NSMAP["w"]
        xml_bytes = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:comments xmlns:w="{w_ns}"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            "/>"
        ).encode("utf-8")

        element = parse_xml(xml_bytes)
        from docx.opc.part import XmlPart

        part = XmlPart(
            PackURI("/word/comments.xml"),
            COMMENTS_CONTENT_TYPE,
            element,
            doc.part.package,
        )
        doc.part.relate_to(part, COMMENTS_REL_TYPE)
        return part

    def _get_next_comment_id(self, doc, comments_part) -> int:
        """Get the next available comment ID by checking comments.xml and body markers."""
        return self._max_markup_id(doc, comments_part) + 1

    def _add_comment_xml(self, comments_part, comment_id: int, text: str, author: str):
        """Add a comment element to the comments XML part's element tree."""
        root = comments_part.element

        date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        initials = "".join(word[0].upper() for word in author.split() if word)

        comment_el = OxmlElement("w:comment")
        comment_el.set(qn("w:id"), str(comment_id))
        comment_el.set(qn("w:author"), author)
        comment_el.set(qn("w:date"), date)
        comment_el.set(qn("w:initials"), initials)

        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        t_el = OxmlElement("w:t")
        t_el.text = text
        r_el.append(t_el)
        p_el.append(r_el)
        comment_el.append(p_el)
        root.append(comment_el)

    def _insert_comment_markers(self, start_run_el, end_run_el, comment_id: int):
        """Insert commentRangeStart, commentRangeEnd, and commentReference.

        Places markers outside any existing comment markers so overlapping
        comments produce flat sibling sequences rather than nested ranges.
        """
        range_start = OxmlElement("w:commentRangeStart")
        range_start.set(qn("w:id"), str(comment_id))

        range_end = OxmlElement("w:commentRangeEnd")
        range_end.set(qn("w:id"), str(comment_id))

        ref_run = OxmlElement("w:r")
        ref_rpr = OxmlElement("w:rPr")
        ref_style = OxmlElement("w:rStyle")
        ref_style.set(qn("w:val"), "CommentReference")
        ref_rpr.append(ref_style)
        ref_run.append(ref_rpr)
        ref_el = OxmlElement("w:commentReference")
        ref_el.set(qn("w:id"), str(comment_id))
        ref_run.append(ref_el)

        insert_before = start_run_el
        while True:
            prev = insert_before.getprevious()
            if prev is None:
                break
            tag = _local_tag(prev)
            if tag != "commentRangeStart":
                break
            insert_before = prev
        insert_before.addprevious(range_start)

        insert_after = end_run_el
        while True:
            nxt = insert_after.getnext()
            if nxt is None:
                break
            tag = _local_tag(nxt)
            if tag not in ("commentRangeEnd", "r"):
                break
            if tag == "r":
                has_ref = nxt.find(qn("w:commentReference")) is not None
                if not has_ref:
                    break
            insert_after = nxt
        insert_after.addnext(ref_run)
        ref_run.addprevious(range_end)

    def _extract_comment_contexts(self, body: ET.Element) -> dict:
        """Extract the text between commentRangeStart and commentRangeEnd markers.

        Args:
            body: The document body XML element

        Returns:
            Dict mapping comment ID to context text
        """
        # Flatten all elements in document order
        all_elements = list(body.iter())

        # Find comment range markers
        range_starts = {}
        range_ends = {}

        for i, el in enumerate(all_elements):
            tag = _local_tag(el)
            if tag == "commentRangeStart":
                cid = el.get(f'{{{NSMAP["w"]}}}id')
                if cid:
                    range_starts[cid] = i
            elif tag == "commentRangeEnd":
                cid = el.get(f'{{{NSMAP["w"]}}}id')
                if cid:
                    range_ends[cid] = i

        # Extract text between start and end markers
        context_map = {}
        for cid, start_idx in range_starts.items():
            end_idx = range_ends.get(cid)
            if end_idx is None:
                continue

            texts = []
            for el in all_elements[start_idx:end_idx + 1]:
                tag = _local_tag(el)
                if tag == "t" and el.text:
                    texts.append(el.text)

            context = "".join(texts).strip()
            if context:
                context_map[cid] = context

        return context_map


_client = None


def get_client() -> MswordClient:
    """Get or create the global MswordClient instance."""
    global _client
    if _client is None:
        _client = MswordClient()
    return _client
